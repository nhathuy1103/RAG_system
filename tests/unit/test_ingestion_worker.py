"""Unit tests for the durable Storage-to-Advanced-Extraction worker."""

from collections.abc import Mapping, Sequence
from dataclasses import replace
from hashlib import sha256
from io import BytesIO
from types import SimpleNamespace
from typing import NoReturn
from uuid import NAMESPACE_URL, UUID, uuid5

import pytest
from docx import Document as DocxDocument

from app.bootstrap.settings import Settings as AppSettings
from app.ingestion.application.worker import (
    IngestionWorker,
    build_ingestion_profile,
)
from app.ingestion.domain.models import (
    ClaimedIngestionJob,
    PersistedChunk,
)
from app.ingestion.ports.repositories import IngestionRepositoryError
from app.knowledge_quality.application.analysis import (
    build_legacy_document_fingerprint,
)
from app.knowledge_quality.domain.models import (
    ChunkDedupCandidate,
    DocumentFingerprint,
)
from app.pipeline.bootstrap.composition import build_ingestion_embedding_pipeline
from app.pipeline.bootstrap.settings import Settings
from app.pipeline.documents.adapters.parsers import CsvParser, TxtParser
from app.pipeline.documents.application.content_identity import (
    build_parsed_document_fingerprint,
)
from app.pipeline.indexing.adapters.vector_indexes import InMemoryVectorIndex
from app.pipeline.indexing.domain.context_enrichment import (
    ChunkContextEnrichment,
    ChunkContextEnrichmentRequest,
)
from app.pipeline.indexing.domain.embedded_chunk import EmbeddedChunk
from app.structured_facts.application.persistence import (
    build_structured_fact_persistence_batch,
)
from app.structured_facts.application.table_analyzer import analyze_table
from app.structured_facts.ports.repositories import (
    StructuredClaimCandidate,
    StructuredFactWriteResult,
)

OWNER_ID = UUID("20000000-0000-0000-0000-000000000002")
NOTEBOOK_ID = UUID("10000000-0000-0000-0000-000000000001")
DOCUMENT_ID = UUID("30000000-0000-0000-0000-000000000003")
JOB_ID = UUID("40000000-0000-0000-0000-000000000004")
CONTENT = b"Revenue increased in 2026.\n\nCosts stayed controlled."
CSV_CONTENT = (
    "Dự án,Tòa,Mã căn,Giá bán,Ngày hiệu lực\nSunrise,A,U01,3000000000,2025-03-01\n"
).encode()


def make_job(
    *,
    content_hash: str | None = None,
    knowledge_quality_mode: str = "off",
    structured_fact_mode: str = "off",
) -> ClaimedIngestionJob:
    return ClaimedIngestionJob(
        id=JOB_ID,
        owner_id=OWNER_ID,
        notebook_id=NOTEBOOK_ID,
        document_id=DOCUMENT_ID,
        attempt_number=1,
        configuration={
            "knowledge_quality_mode": knowledge_quality_mode,
            "structured_fact_mode": structured_fact_mode,
        },
        storage_bucket="documents",
        storage_object_path=f"{OWNER_ID}/{NOTEBOOK_ID}/{DOCUMENT_ID}/brief.txt",
        original_filename="brief.txt",
        mime_type="text/plain",
        size_bytes=len(CONTENT),
        content_hash=content_hash or sha256(CONTENT).hexdigest(),
    )


def make_csv_job(*, structured_fact_mode: str = "shadow") -> ClaimedIngestionJob:
    return replace(
        make_job(structured_fact_mode=structured_fact_mode),
        storage_object_path=f"{OWNER_ID}/{NOTEBOOK_ID}/{DOCUMENT_ID}/prices.csv",
        original_filename="prices.csv",
        mime_type="text/csv",
        size_bytes=len(CSV_CONTENT),
        content_hash=sha256(CSV_CONTENT).hexdigest(),
    )


def make_prior_csv_candidates() -> tuple[StructuredClaimCandidate, ...]:
    table = CsvParser().parse(CSV_CONTENT).tables[0]
    analysis = analyze_table(document_id=str(UUID(int=9)), table=table)
    batch = build_structured_fact_persistence_batch(
        analyses=(analysis,),
        tables=(table,),
        embedded_chunks=(),
    )
    snapshot = batch.table_snapshots[0]
    snapshot_id = UUID(int=10)
    return tuple(
        StructuredClaimCandidate(
            claim_id=uuid5(NAMESPACE_URL, f"prior-claim:{claim.claim_identity_hash}"),
            snapshot_id=snapshot_id,
            document_id=UUID(analysis.document_id),
            document_version=1,
            snapshot_key=str(snapshot["snapshot_key"]),
            schema_fingerprint=str(snapshot["schema_fingerprint"]),
            template_fingerprint=None,
            normalized_schema=snapshot["normalized_schema"],
            candidate_identity_hash=claim.candidate_identity_hash,
            claim=claim.to_payload(),
        )
        for claim in analysis.claims
    )


class FakeWorkerRepository:
    def __init__(
        self,
        job: ClaimedIngestionJob,
        *,
        fail_complete: bool = False,
        duplicate_id: UUID | None = None,
        renew_succeeds: bool = True,
        failure_accepted: bool = True,
        completion_disposition: str = "completed",
    ) -> None:
        self.job = job
        self.fail_complete = fail_complete
        self.duplicate_id = duplicate_id
        self.renew_succeeds = renew_succeeds
        self.failure_accepted = failure_accepted
        self.completion_disposition = completion_disposition
        self.claimed = False
        self.completed: tuple[ClaimedIngestionJob, tuple[PersistedChunk, ...], str, int] | None = (
            None
        )
        self.failed: tuple[ClaimedIngestionJob, str] | None = None
        self.duplicate_completed: tuple[ClaimedIngestionJob, UUID] | None = None
        self.completed_fingerprint = None
        self.completed_relations = ()
        self.completed_quality_mode = "off"

    async def claim(
        self,
        worker_id: str,
        lease_seconds: int,
    ) -> ClaimedIngestionJob | None:
        assert worker_id == "test-worker"
        assert lease_seconds == 60
        if self.claimed:
            return None
        self.claimed = True
        return self.job

    async def renew_lease(
        self,
        job_id: UUID,
        worker_id: str,
        claim_token: UUID,
        lease_seconds: int,
    ) -> bool:
        assert claim_token == self.job.claim_token
        del job_id, worker_id, lease_seconds
        return self.renew_succeeds

    async def complete(
        self,
        job: ClaimedIngestionJob,
        worker_id: str,
        chunks: tuple[PersistedChunk, ...],
        embedding_model: str,
        embedding_dimensions: int,
        fingerprint=None,
        relations=(),
        effective_quality_mode="off",
        claim_scope=None,
    ) -> str:
        assert worker_id == "test-worker"
        if self.fail_complete:
            raise IngestionRepositoryError("job was cancelled")
        self.completed_fingerprint = fingerprint
        self.completed_relations = tuple(relations)
        self.completed_quality_mode = effective_quality_mode
        assert claim_scope is not None
        self.completed = (job, tuple(chunks), embedding_model, embedding_dimensions)
        return self.completion_disposition

    async def fail(
        self,
        job: ClaimedIngestionJob,
        worker_id: str,
        error_message: str,
    ) -> bool:
        assert worker_id == "test-worker"
        self.failed = (job, error_message)
        return self.failure_accepted

    async def find_content_duplicate(self, job, fingerprint):
        del job, fingerprint
        return self.duplicate_id

    async def find_chunk_dedup_candidates(
        self,
        job,
        probes,
        embedding_model,
        candidates_per_probe,
    ):
        del job, probes, embedding_model, candidates_per_probe
        return ()

    async def complete_duplicate(
        self,
        job,
        worker_id,
        canonical_document_id,
        fingerprint,
        effective_quality_mode="off",
    ):
        del fingerprint
        assert worker_id == "test-worker"
        assert effective_quality_mode == "on"
        self.duplicate_completed = (job, canonical_document_id)


class FakeStorage:
    async def upload(
        self,
        bucket: str,
        object_path: str,
        content: bytes,
        mime_type: str,
    ) -> None:
        del bucket, object_path, content, mime_type
        raise NotImplementedError

    async def delete(self, bucket: str, object_path: str) -> None:
        del bucket, object_path
        raise NotImplementedError

    async def download(self, bucket: str, object_path: str) -> bytes:
        assert bucket == "documents"
        assert object_path.endswith("/brief.txt")
        return CONTENT


class PayloadStorage(FakeStorage):
    def __init__(self, content: bytes, expected_suffix: str) -> None:
        self.content = content
        self.expected_suffix = expected_suffix

    async def download(self, bucket: str, object_path: str) -> bytes:
        assert bucket == "documents"
        assert object_path.endswith(self.expected_suffix)
        return self.content


class FakeStructuredFactStore:
    def __init__(
        self,
        *,
        candidates: Sequence[StructuredClaimCandidate] = (),
        candidate_batches: Sequence[Sequence[StructuredClaimCandidate]] | None = None,
        fail_load: bool = False,
        fail_replace_attempts: int = 0,
    ) -> None:
        self.candidates = tuple(candidates)
        self.candidate_batches = (
            tuple(tuple(batch) for batch in candidate_batches)
            if candidate_batches is not None
            else None
        )
        self.fail_load = fail_load
        self.fail_replace_attempts = fail_replace_attempts
        self.load_calls: list[dict[str, object]] = []
        self.replace_calls: list[dict[str, object]] = []

    async def load_claim_candidates(
        self,
        *,
        notebook_id: UUID,
        document_id: UUID,
        candidate_identity_hashes: Sequence[str],
        schema_fingerprints: Sequence[str] = (),
        limit: int = 10000,
    ) -> tuple[StructuredClaimCandidate, ...]:
        call_index = len(self.load_calls)
        self.load_calls.append(
            {
                "notebook_id": notebook_id,
                "document_id": document_id,
                "candidate_identity_hashes": tuple(candidate_identity_hashes),
                "schema_fingerprints": tuple(schema_fingerprints),
                "limit": limit,
            }
        )
        if self.fail_load:
            raise RuntimeError("candidate lookup unavailable")
        if self.candidate_batches is not None:
            return self.candidate_batches[min(call_index, len(self.candidate_batches) - 1)]
        return self.candidates

    async def replace_for_document(
        self,
        *,
        job_id: UUID,
        document_id: UUID,
        extractor_version: str,
        table_snapshots: Sequence[Mapping[str, object]],
        claims: Sequence[Mapping[str, object]],
        relations: Sequence[Mapping[str, object]] = (),
    ) -> StructuredFactWriteResult:
        self.replace_calls.append(
            {
                "job_id": job_id,
                "document_id": document_id,
                "extractor_version": extractor_version,
                "table_snapshots": tuple(table_snapshots),
                "claims": tuple(claims),
                "relations": tuple(relations),
            }
        )
        if self.fail_replace_attempts > 0:
            self.fail_replace_attempts -= 1
            raise RuntimeError("structured replacement unavailable")
        return StructuredFactWriteResult(
            table_count=len(table_snapshots),
            claim_count=len(claims),
            relation_count=len(relations),
        )


class TransactionalVectorIndex(InMemoryVectorIndex):
    """Test double for pgvector writes committed inside the completion RPC."""

    completion_is_transactional = True


def make_worker(
    repository: FakeWorkerRepository,
    *,
    vector_index: InMemoryVectorIndex | None = None,
    knowledge_quality_mode: str = "off",
    structured_fact_mode: str = "off",
    structured_fact_store: FakeStructuredFactStore | None = None,
    object_storage: FakeStorage | None = None,
) -> IngestionWorker:
    settings = Settings(app_env="test")
    pipeline = build_ingestion_embedding_pipeline(
        settings,
        vector_index=vector_index or InMemoryVectorIndex(),
    )
    return IngestionWorker(
        repository=repository,
        object_storage=object_storage or FakeStorage(),
        pipeline=pipeline,
        poll_interval_seconds=0.01,
        lease_seconds=60,
        worker_id="test-worker",
        knowledge_quality_mode=knowledge_quality_mode,
        structured_fact_mode=structured_fact_mode,
        structured_fact_store=structured_fact_store,
    )


@pytest.mark.anyio
async def test_worker_runs_advanced_extraction_and_persists_canonical_chunks() -> None:
    repository = FakeWorkerRepository(make_job())

    processed = await make_worker(repository).run_once()

    assert processed is True
    assert repository.failed is None
    assert repository.completed is not None
    job, chunks, model, dimensions = repository.completed
    assert job.id == JOB_ID
    assert chunks
    assert chunks[0].content.startswith("Revenue increased")
    assert chunks[0].token_count > 0
    assert chunks[0].metadata["qdrant_point_id"] == str(chunks[0].id)
    assert chunks[0].metadata["parser_name"] == "txt"
    assert "token_count" not in chunks[0].metadata
    assert isinstance(chunks[0].metadata["source_block_ids"], list)
    assert chunks[0].metadata["page_number"] == 1
    assert "canonical_text" not in chunks[0].metadata
    assert "provenance_metadata" not in chunks[0].metadata
    assert "authority_metadata" not in chunks[0].metadata
    assert len(chunks[0].embedding) == 32
    assert chunks[0].metadata["normalized_content_hash"]
    assert chunks[0].metadata["loose_content_signature"]
    assert chunks[0].metadata["normalization_version"]
    assert model == "local-hash-embedding-v1"
    assert dimensions == 32


def test_persisted_chunk_keeps_distinct_canonical_and_nonempty_lineage_metadata() -> None:
    chunk = EmbeddedChunk(
        id="source-chunk",
        document_id=str(DOCUMENT_ID),
        document_version=1,
        owner_id=str(OWNER_ID),
        tenant_id=str(NOTEBOOK_ID),
        chunk_index=0,
        page_number=1,
        section_title="Summary",
        checksum="checksum",
        text="Normalized chunk text",
        canonical_text="Canonical source text",
        token_count=3,
        embedding=(0.1, 0.2),
        embedding_model="test-model",
        provenance_metadata={"source": "pdf"},
        authority_metadata={"authority": "approved-policy"},
    )
    result = SimpleNamespace(embedded_chunks=(chunk,))

    persisted = IngestionWorker._to_persisted_chunks(result)[0]  # type: ignore[arg-type]

    assert persisted.metadata["canonical_text"] == "Canonical source text"
    assert persisted.metadata["provenance_metadata"] == {"source": "pdf"}
    assert persisted.metadata["authority_metadata"] == {"authority": "approved-policy"}


@pytest.mark.anyio
async def test_transactional_vector_backend_persists_embeddings_only_via_completion() -> None:
    repository = FakeWorkerRepository(make_job())
    vector_index = TransactionalVectorIndex()

    processed = await make_worker(
        repository,
        vector_index=vector_index,
    ).run_once()

    assert processed is True
    assert repository.completed is not None
    assert repository.completed[1][0].embedding
    assert vector_index.list_chunks() == []


@pytest.mark.anyio
async def test_worker_fails_closed_when_storage_checksum_does_not_match() -> None:
    repository = FakeWorkerRepository(make_job(content_hash="0" * 64))

    processed = await make_worker(repository).run_once()

    assert processed is True
    assert repository.completed is None
    assert repository.failed is not None
    assert "checksum does not match" in repository.failed[1]


@pytest.mark.anyio
async def test_worker_removes_vectors_when_deleted_document_cannot_complete() -> None:
    repository = FakeWorkerRepository(make_job(), fail_complete=True)
    vector_index = InMemoryVectorIndex()

    processed = await make_worker(
        repository,
        vector_index=vector_index,
    ).run_once()

    assert processed is True
    assert repository.completed is None
    assert repository.failed is not None
    assert vector_index.list_chunks() == []


@pytest.mark.anyio
async def test_accepted_failure_deletes_only_its_vector_generation() -> None:
    repository = FakeWorkerRepository(make_job(), fail_complete=True)
    vector_index = InMemoryVectorIndex()
    vector_index.upsert_chunks(
        [
            EmbeddedChunk(
                id="new-generation-chunk",
                document_id=str(DOCUMENT_ID),
                document_version=1,
                owner_id=str(OWNER_ID),
                tenant_id=str(NOTEBOOK_ID),
                chunk_index=99,
                page_number=1,
                section_title=None,
                checksum="new-generation",
                text="A newer generation remains available.",
                canonical_text="A newer generation remains available.",
                token_count=6,
                embedding=(1.0, 0.0),
                embedding_model="local-hash-embedding-v1",
                metadata={"ingestion_generation": "new-generation"},
            )
        ]
    )

    processed = await make_worker(
        repository,
        vector_index=vector_index,
    ).run_once()

    assert processed is True
    assert [chunk.id for chunk in vector_index.list_chunks()] == ["new-generation-chunk"]


@pytest.mark.anyio
async def test_worker_auto_aliases_strict_content_duplicate_without_embedding() -> None:
    canonical_id = UUID("50000000-0000-0000-0000-000000000005")
    job = replace(
        make_job(knowledge_quality_mode="on"),
        configuration={
            "knowledge_quality_mode": "on",
            "contextual_enrichment_enabled": True,
        },
    )
    repository = FakeWorkerRepository(
        job,
        duplicate_id=canonical_id,
    )
    vector_index = InMemoryVectorIndex()

    class FailIfContextualized:
        document_context_char_limit = 1000
        profile: dict[str, object] = {}
        calls = 0

        def enrich(self, request: object) -> NoReturn:
            del request
            self.calls += 1
            raise AssertionError("Exact document duplicate must not incur LLM context cost")

    context_enricher = FailIfContextualized()
    worker = make_worker(
        repository,
        vector_index=vector_index,
        knowledge_quality_mode="on",
    )
    worker._pipeline.context_enricher = context_enricher

    processed = await worker.run_once()

    assert processed is True
    assert repository.completed is None
    assert repository.failed is None
    assert repository.duplicate_completed == (repository.job, canonical_id)
    assert vector_index.list_chunks() == []
    assert context_enricher.calls == 0


@pytest.mark.anyio
async def test_worker_persists_generated_context_for_new_profile() -> None:
    job = replace(
        make_job(),
        configuration={
            "knowledge_quality_mode": "off",
            "contextual_enrichment_enabled": True,
            "contextual_enrichment_model": "context-test-v1",
            "contextual_enrichment_prompt_version": "chunk-context-test-v1",
            "contextual_text_version": "contextual-text-v4",
        },
    )
    repository = FakeWorkerRepository(job)

    class ContextEnricher:
        document_context_char_limit = 2000
        profile = {
            "contextual_enrichment_model": "context-test-v1",
            "contextual_enrichment_prompt_version": "chunk-context-test-v1",
        }

        def enrich(
            self,
            request: ChunkContextEnrichmentRequest,
        ) -> ChunkContextEnrichment:
            assert request.chunk_text
            return ChunkContextEnrichment(
                context_text="This chunk summarizes the 2026 revenue statement.",
                search_terms=("Revenue", "2026"),
                status="generated",
                provider="test",
                model="context-test-v1",
                prompt_version="chunk-context-test-v1",
                input_checksum="a" * 64,
                source_scope=request.source_scope,
            )

    worker = make_worker(repository)
    worker._pipeline.context_enricher = ContextEnricher()

    processed = await worker.run_once()

    assert processed is True
    assert repository.completed is not None
    persisted = repository.completed[1][0]
    assert persisted.content in CONTENT.decode("utf-8")
    retrieval = persisted.metadata["retrieval_metadata"]
    assert isinstance(retrieval, dict)
    assert retrieval["contextual_summary"] == ("This chunk summarizes the 2026 revenue statement.")
    assert "contextual_search_terms" not in retrieval
    enrichment = persisted.metadata["context_enrichment"]
    assert isinstance(enrichment, dict)
    assert enrichment["status"] == "generated"
    assert persisted.metadata["contextual_text_version"] == "contextual-text-v4"


@pytest.mark.anyio
async def test_worker_auto_aliases_same_content_from_txt_and_docx_projection() -> None:
    canonical_id = UUID("50000000-0000-0000-0000-000000000005")
    canonical_fingerprint = build_parsed_document_fingerprint(TxtParser().parse(CONTENT))
    document = DocxDocument()
    for paragraph in CONTENT.decode("utf-8").split("\n\n"):
        document.add_paragraph(paragraph)
    output = BytesIO()
    document.save(output)
    docx_content = output.getvalue()
    docx_job = replace(
        make_job(knowledge_quality_mode="on"),
        original_filename="brief.docx",
        mime_type=("application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        size_bytes=len(docx_content),
        content_hash=sha256(docx_content).hexdigest(),
        storage_object_path=(f"{OWNER_ID}/{NOTEBOOK_ID}/{DOCUMENT_ID}/brief.docx"),
    )

    class CrossFormatRepository(FakeWorkerRepository):
        async def find_content_duplicate(
            self,
            job: ClaimedIngestionJob,
            fingerprint: DocumentFingerprint,
        ) -> UUID | None:
            assert job == docx_job
            assert fingerprint.strict_hash == canonical_fingerprint.strict_hash
            assert fingerprint.normalization_version == canonical_fingerprint.normalization_version
            assert fingerprint.loose_signature == canonical_fingerprint.loose_signature
            return canonical_id

    class DocxStorage(FakeStorage):
        async def download(self, bucket: str, object_path: str) -> bytes:
            assert bucket == "documents"
            assert object_path.endswith("/brief.docx")
            return docx_content

    repository = CrossFormatRepository(docx_job)
    worker = make_worker(
        repository,
        vector_index=InMemoryVectorIndex(),
        knowledge_quality_mode="on",
        object_storage=DocxStorage(),
    )

    processed = await worker.run_once()

    assert processed is True
    assert repository.duplicate_completed == (docx_job, canonical_id)
    assert repository.completed is None


@pytest.mark.anyio
async def test_shadow_job_persists_candidate_without_automatic_aliasing() -> None:
    canonical_id = UUID("50000000-0000-0000-0000-000000000005")
    repository = FakeWorkerRepository(
        make_job(knowledge_quality_mode="shadow"),
        duplicate_id=canonical_id,
    )

    processed = await make_worker(
        repository,
        # A changed worker flag must not upgrade queued shadow work to on.
        knowledge_quality_mode="on",
    ).run_once()

    assert processed is True
    assert repository.failed is None
    assert repository.duplicate_completed is None
    assert repository.completed is not None
    assert repository.completed_fingerprint is not None
    assert repository.completed_quality_mode == "shadow"
    assert len(repository.completed_relations) == 1
    assert repository.completed_relations[0].target_document_id == canonical_id
    assert repository.completed_relations[0].relation_type.value == "exact_content"


@pytest.mark.anyio
async def test_worker_reuses_exact_chunk_embedding_before_provider_call() -> None:
    reused_vector = (0.25,) * 32
    target_document_id = UUID("50000000-0000-0000-0000-000000000005")

    class ExactChunkRepository(FakeWorkerRepository):
        async def find_chunk_dedup_candidates(
            self,
            job,
            probes,
            embedding_model,
            candidates_per_probe,
        ):
            assert job == self.job
            assert embedding_model == "local-hash-embedding-v1"
            assert candidates_per_probe == 5
            probe = probes[0]
            return (
                ChunkDedupCandidate(
                    source_chunk_index=probe.chunk_index,
                    target_chunk_id="60000000-0000-0000-0000-000000000006",
                    target_document_id=target_document_id,
                    target_chunk_index=2,
                    canonical_text=probe.canonical_text,
                    normalized_content_hash=probe.fingerprint.strict_hash,
                    normalization_version=probe.fingerprint.normalization_version,
                    loose_content_signature=probe.fingerprint.loose_signature,
                    embedding_text_checksum=probe.embedding_text_checksum,
                    embedding=reused_vector,
                    embedding_model=embedding_model,
                    lsh_band_matches=8,
                ),
            )

    repository = ExactChunkRepository(make_job(knowledge_quality_mode="on"))

    processed = await make_worker(
        repository,
        knowledge_quality_mode="on",
    ).run_once()

    assert processed is True
    assert repository.failed is None
    assert repository.completed is not None
    persisted = repository.completed[1][0]
    assert persisted.embedding == reused_vector
    quality = persisted.metadata["pre_embedding_quality"]
    assert isinstance(quality, dict)
    assert quality["action"] == "reuse_exact_embedding"
    assert quality["embedding_reused"] is True


@pytest.mark.anyio
async def test_runtime_shadow_is_a_kill_switch_for_queued_on_job() -> None:
    canonical_id = UUID("50000000-0000-0000-0000-000000000005")
    repository = FakeWorkerRepository(
        make_job(knowledge_quality_mode="on"),
        duplicate_id=canonical_id,
    )

    processed = await make_worker(
        repository,
        knowledge_quality_mode="shadow",
    ).run_once()

    assert processed is True
    assert repository.duplicate_completed is None
    assert repository.completed is not None
    assert repository.completed_quality_mode == "shadow"
    assert repository.completed_fingerprint is not None
    assert len(repository.completed_relations) == 1


@pytest.mark.anyio
async def test_repair_skips_quality_decisions_but_checks_existing_fingerprint() -> None:
    expected = build_parsed_document_fingerprint(TxtParser().parse(CONTENT))
    job = replace(
        make_job(knowledge_quality_mode="on"),
        configuration={
            "knowledge_quality_mode": "on",
            "ingestion_kind": "reconciliation_repair",
            "expected_normalized_content_hash": expected.strict_hash,
            "expected_normalization_version": expected.normalization_version,
            "expected_loose_content_signature": expected.loose_signature,
        },
    )
    repository = FakeWorkerRepository(
        job,
        duplicate_id=UUID("50000000-0000-0000-0000-000000000005"),
    )

    processed = await make_worker(
        repository,
        knowledge_quality_mode="on",
    ).run_once()

    assert processed is True
    assert repository.failed is None
    assert repository.duplicate_completed is None
    assert repository.completed is not None
    assert repository.completed_fingerprint == expected
    assert repository.completed_relations == ()
    assert repository.completed_quality_mode == "off"


@pytest.mark.anyio
async def test_repair_can_verify_legacy_v1_document_fingerprint() -> None:
    expected = build_legacy_document_fingerprint(CONTENT.decode("utf-8"))
    job = replace(
        make_job(knowledge_quality_mode="on"),
        configuration={
            "knowledge_quality_mode": "on",
            "ingestion_kind": "reconciliation_repair",
            "expected_normalized_content_hash": expected.strict_hash,
            "expected_normalization_version": expected.normalization_version,
            "expected_loose_content_signature": expected.loose_signature,
        },
    )
    repository = FakeWorkerRepository(job)

    processed = await make_worker(
        repository,
        knowledge_quality_mode="on",
    ).run_once()

    assert processed is True
    assert repository.failed is None
    assert repository.completed_fingerprint == expected
    assert repository.completed_quality_mode == "off"


@pytest.mark.anyio
async def test_repair_fails_before_commit_when_source_fingerprint_changed() -> None:
    expected = build_parsed_document_fingerprint(TxtParser().parse(CONTENT))
    job = replace(
        make_job(knowledge_quality_mode="on"),
        configuration={
            "knowledge_quality_mode": "on",
            "ingestion_kind": "reconciliation_repair",
            "expected_normalized_content_hash": "f" * 64,
            "expected_normalization_version": expected.normalization_version,
            "expected_loose_content_signature": expected.loose_signature,
        },
    )
    repository = FakeWorkerRepository(job)

    processed = await make_worker(
        repository,
        knowledge_quality_mode="on",
    ).run_once()

    assert processed is True
    assert repository.completed is None
    assert repository.failed is not None
    assert "fingerprint changed" in repository.failed[1]


@pytest.mark.anyio
async def test_atomic_duplicate_suppression_removes_external_generation() -> None:
    repository = FakeWorkerRepository(
        make_job(knowledge_quality_mode="on"),
        completion_disposition="duplicate_suppressed",
    )
    vector_index = InMemoryVectorIndex()

    processed = await make_worker(
        repository,
        vector_index=vector_index,
        knowledge_quality_mode="on",
    ).run_once()

    assert processed is True
    assert repository.completed is not None
    assert repository.failed is None
    assert vector_index.list_chunks() == []


@pytest.mark.anyio
async def test_legacy_job_without_durable_mode_fails_safe_to_off() -> None:
    canonical_id = UUID("50000000-0000-0000-0000-000000000005")
    job = replace(
        make_job(knowledge_quality_mode="on"),
        configuration={},
    )
    repository = FakeWorkerRepository(job, duplicate_id=canonical_id)

    processed = await make_worker(
        repository,
        knowledge_quality_mode="on",
    ).run_once()

    assert processed is True
    assert repository.duplicate_completed is None
    assert repository.completed is not None
    assert repository.completed_fingerprint is None
    assert repository.completed_relations == ()
    assert repository.completed_quality_mode == "off"


@pytest.mark.anyio
async def test_worker_does_not_write_vectors_after_generation_lease_is_lost() -> None:
    repository = FakeWorkerRepository(make_job(), renew_succeeds=False)
    vector_index = InMemoryVectorIndex()

    processed = await make_worker(
        repository,
        vector_index=vector_index,
    ).run_once()

    assert processed is True
    assert repository.completed is None
    assert repository.failed is not None
    assert vector_index.list_chunks() == []


@pytest.mark.anyio
async def test_stale_failure_cannot_delete_vectors_owned_by_new_generation() -> None:
    repository = FakeWorkerRepository(
        make_job(),
        fail_complete=True,
        failure_accepted=False,
    )
    vector_index = InMemoryVectorIndex()

    processed = await make_worker(
        repository,
        vector_index=vector_index,
    ).run_once()

    assert processed is True
    assert repository.completed is None
    assert repository.failed is not None
    assert vector_index.list_chunks()


def test_ingestion_profile_requires_openai_qdrant_contract() -> None:
    profile = build_ingestion_profile(
        Settings(
            app_env="test",
            embedding_provider="openai",
            openai_api_key="test-key",
            openai_embedding_model="text-embedding-3-small",
            vector_store_backend="qdrant",
        )
    )

    assert profile.embedding_model == "text-embedding-3-small"
    assert profile.embedding_dimensions == 1536
    assert profile.configuration["advanced_extraction_enabled"] is True
    assert profile.configuration["chunking_strategy"] == "structure_recursive"
    assert profile.configuration["contextual_enrichment_enabled"] is False
    assert profile.configuration["contextual_enrichment_model"] == "gpt-4o-mini"
    assert profile.configuration["contextual_enrichment_prompt_version"] == "chunk-context-v4"
    assert profile.configuration["contextual_text_version"] == "contextual-text-v4"
    assert profile.configuration["embedding_provider"] == "openai"
    assert profile.configuration["vector_store_backend"] == "qdrant"
    assert profile.configuration["knowledge_quality_mode"] == "off"
    assert profile.configuration["structured_fact_mode"] == "off"


def test_application_default_enables_safe_exact_content_reuse() -> None:
    assert AppSettings.model_validate({}).knowledge_quality_mode == "on"
    assert AppSettings.model_validate({}).structured_fact_mode == "off"


def test_ingestion_profile_rejects_silent_local_fallback() -> None:
    with pytest.raises(ValueError, match="EMBEDDING_PROVIDER=openai"):
        build_ingestion_profile(Settings(app_env="test"))


def test_ingestion_profile_persists_rollout_mode() -> None:
    profile = build_ingestion_profile(
        Settings(
            app_env="test",
            embedding_provider="openai",
            openai_api_key="test-key",
            openai_embedding_model="text-embedding-3-small",
            vector_store_backend="pgvector",
        ),
        knowledge_quality_mode="on",
        structured_fact_mode="shadow",
    )

    assert profile.configuration["knowledge_quality_mode"] == "on"
    assert profile.configuration["structured_fact_mode"] == "shadow"


def test_structured_fact_mode_uses_safer_durable_runtime_value() -> None:
    repository = FakeWorkerRepository(make_job())
    worker = make_worker(repository)
    shadow_job = replace(
        make_job(),
        configuration={
            "knowledge_quality_mode": "off",
            "structured_fact_mode": "shadow",
        },
    )

    assert worker._structured_mode_for_job(shadow_job) == "off"


@pytest.mark.anyio
async def test_structured_shadow_persists_csv_facts_after_document_completion() -> None:
    repository = FakeWorkerRepository(make_csv_job())
    fact_store = FakeStructuredFactStore()

    processed = await make_worker(
        repository,
        structured_fact_mode="shadow",
        structured_fact_store=fact_store,
        object_storage=PayloadStorage(CSV_CONTENT, "/prices.csv"),
    ).run_once()

    assert processed is True
    assert repository.completed is not None
    assert repository.failed is None
    assert repository.completed_fingerprint is None
    assert len(fact_store.load_calls) == 2
    assert len(fact_store.replace_calls) == 1
    replacement = fact_store.replace_calls[0]
    assert len(replacement["table_snapshots"]) == 1
    assert replacement["claims"]
    assert replacement["relations"] == ()
    assert fact_store.load_calls[0]["candidate_identity_hashes"]
    assert fact_store.load_calls[0]["schema_fingerprints"]


@pytest.mark.anyio
async def test_structured_worker_diffs_loaded_candidate_before_atomic_replace() -> None:
    repository = FakeWorkerRepository(make_csv_job())
    fact_store = FakeStructuredFactStore(candidates=make_prior_csv_candidates())

    processed = await make_worker(
        repository,
        structured_fact_mode="shadow",
        structured_fact_store=fact_store,
        object_storage=PayloadStorage(CSV_CONTENT, "/prices.csv"),
    ).run_once()

    assert processed is True
    assert repository.completed is not None
    assert repository.failed is None
    relations = fact_store.replace_calls[0]["relations"]
    assert len(relations) == 1
    assert relations[0]["relation_type"] == "unchanged"
    assert relations[0]["review_status"] == "auto_confirmed"
    assert relations[0]["source_claim_key"]
    assert relations[0]["target_claim_key"]


@pytest.mark.anyio
async def test_structured_worker_reloads_after_write_to_close_concurrent_ingest_race() -> None:
    repository = FakeWorkerRepository(make_csv_job())
    fact_store = FakeStructuredFactStore(candidate_batches=((), make_prior_csv_candidates()))

    processed = await make_worker(
        repository,
        structured_fact_mode="shadow",
        structured_fact_store=fact_store,
        object_storage=PayloadStorage(CSV_CONTENT, "/prices.csv"),
    ).run_once()

    assert processed is True
    assert repository.completed is not None
    assert len(fact_store.load_calls) == 2
    assert len(fact_store.replace_calls) == 2
    assert fact_store.replace_calls[0]["relations"] == ()
    assert fact_store.replace_calls[1]["relations"][0]["relation_type"] == "unchanged"


@pytest.mark.anyio
async def test_structured_runtime_off_is_kill_switch_for_queued_shadow_job() -> None:
    repository = FakeWorkerRepository(make_csv_job())
    fact_store = FakeStructuredFactStore()

    processed = await make_worker(
        repository,
        structured_fact_mode="off",
        structured_fact_store=fact_store,
        object_storage=PayloadStorage(CSV_CONTENT, "/prices.csv"),
    ).run_once()

    assert processed is True
    assert repository.completed is not None
    assert fact_store.load_calls == []
    assert fact_store.replace_calls == []


@pytest.mark.anyio
async def test_structured_candidate_failure_still_replaces_current_facts() -> None:
    repository = FakeWorkerRepository(make_csv_job())
    fact_store = FakeStructuredFactStore(fail_load=True)

    processed = await make_worker(
        repository,
        structured_fact_mode="shadow",
        structured_fact_store=fact_store,
        object_storage=PayloadStorage(CSV_CONTENT, "/prices.csv"),
    ).run_once()

    assert processed is True
    assert repository.completed is not None
    assert repository.failed is None
    assert len(fact_store.load_calls) == 2
    assert len(fact_store.replace_calls) == 1
    assert fact_store.replace_calls[0]["relations"] == ()


@pytest.mark.anyio
async def test_structured_replacement_failure_retries_without_failing_ingestion() -> None:
    repository = FakeWorkerRepository(make_csv_job())
    fact_store = FakeStructuredFactStore(fail_replace_attempts=2)

    processed = await make_worker(
        repository,
        structured_fact_mode="shadow",
        structured_fact_store=fact_store,
        object_storage=PayloadStorage(CSV_CONTENT, "/prices.csv"),
    ).run_once()

    assert processed is True
    assert repository.completed is not None
    assert repository.failed is None
    assert len(fact_store.replace_calls) == 2


@pytest.mark.anyio
async def test_structured_mode_clears_stale_facts_when_document_has_no_tables() -> None:
    repository = FakeWorkerRepository(make_job(structured_fact_mode="shadow"))
    fact_store = FakeStructuredFactStore()

    processed = await make_worker(
        repository,
        structured_fact_mode="shadow",
        structured_fact_store=fact_store,
    ).run_once()

    assert processed is True
    assert repository.completed is not None
    assert len(fact_store.replace_calls) == 1
    assert fact_store.replace_calls[0]["table_snapshots"] == ()
    assert fact_store.replace_calls[0]["claims"] == ()
