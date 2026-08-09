"""Build a 300-query metadata benchmark from the uploaded Vinhomes documents.

The benchmark uses the six DOCX files as evidence sources and keeps the three
older pilot documents in the retrieval corpus as distractors. Ground truth is
anchored to deterministic production chunk IDs generated from the source files.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
import unicodedata
from collections import Counter, defaultdict
from collections.abc import Iterable
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document

try:
    from .run_abcd_experiment import EvalChunk, _contains, build_corpus
except ImportError:  # Direct script execution.
    from run_abcd_experiment import EvalChunk, _contains, build_corpus

ROOT = Path(__file__).resolve().parent
DEFAULT_OUTPUT_DIR = ROOT / "real_benchmark_v3"
BASE_GOLD_PATH = ROOT / "gold_metadata.json"

PRICE_DOCUMENTS = {
    2023: "Vinhomes_Gia_Nha_2023.docx",
    2024: "Vinhomes_Gia_Nha_2024.docx",
    2025: "Vinhomes_Gia_Nha_2025.docx",
    2026: "Vinhomes_Gia_Nha_2026.docx",
}
UTILITY_DOCUMENT = "Vinhomes_Tien_Ich_Toan_Quoc.docx"
PLAN_DOCUMENT = "Vinhomes_Ke_Hoach_Xay_Dung_Trien_Khai.docx"
QUERY_SOURCE_DOCUMENTS = (*PRICE_DOCUMENTS.values(), UTILITY_DOCUMENT, PLAN_DOCUMENT)

PRIMARY_SLICES = (
    "content_only",
    "explicit_filter",
    "implicit_filter",
    "cross_document_confusion",
    "version_conflict",
    "section_localization",
    "table_structured",
    "multi_hop",
    "null_insufficient",
    "permission_sensitive",
)

ABLATION_MATRIX = [
    {
        "variant": "v0_raw_text",
        "fields": [],
        "purpose": "Baseline chỉ dùng nội dung chunk.",
    },
    {
        "variant": "v1_document_identity",
        "fields": ["title", "document_type"],
        "purpose": "Đo tác động của định danh tài liệu.",
    },
    {
        "variant": "v2_section_structure",
        "fields": ["section_title", "section_path"],
        "purpose": "Đo khả năng định vị đúng heading.",
    },
    {
        "variant": "v3_block_aware",
        "fields": ["content_kind", "table_header"],
        "purpose": "Đo retrieval trên bảng và loại block.",
    },
    {
        "variant": "v4_context_summary",
        "fields": ["contextual_summary"],
        "purpose": "Đo context ngắn có làm rõ ý nghĩa chunk hay không.",
    },
    {
        "variant": "v5_context_terms",
        "fields": ["contextual_search_terms", "keyword_aliases"],
        "purpose": "Đo từ khóa mở rộng trong hybrid/BM25.",
    },
    {
        "variant": "v6_domain_metadata",
        "fields": [
            "year",
            "document_version",
            "lifecycle_status",
            "effective_status",
            "as_of_date",
            "project_name",
            "project_code",
            "project_status",
            "region",
            "market_type",
            "reliability_grade",
            "source",
            "source_code",
        ],
        "common_base": "v5_context_terms",
        "placement": {
            "structured_filter": True,
            "search_text": True,
            "embedding_text": True,
        },
        "purpose": "Đo metadata chuyên ngành bất động sản và phiên bản.",
    },
    {
        "variant": "v6a_filter_only",
        "fields": [
            "year",
            "document_version",
            "lifecycle_status",
            "effective_status",
            "as_of_date",
            "project_name",
            "project_code",
            "project_status",
            "region",
            "market_type",
            "reliability_grade",
            "source",
            "source_code",
        ],
        "common_base": "v5_context_terms",
        "placement": {
            "structured_filter": True,
            "search_text": False,
            "embedding_text": False,
        },
        "purpose": "Isolate the value of structured domain filters.",
    },
    {
        "variant": "v6b_filter_plus_search_text",
        "fields": [
            "year",
            "document_version",
            "lifecycle_status",
            "effective_status",
            "as_of_date",
            "project_name",
            "project_code",
            "project_status",
            "region",
            "market_type",
            "reliability_grade",
            "source",
            "source_code",
        ],
        "common_base": "v5_context_terms",
        "placement": {
            "structured_filter": True,
            "search_text": True,
            "embedding_text": False,
        },
        "purpose": "Measure the incremental value of domain metadata in lexical search text.",
    },
    {
        "variant": "v6c_filter_plus_embedding_text",
        "fields": [
            "year",
            "document_version",
            "lifecycle_status",
            "effective_status",
            "as_of_date",
            "project_name",
            "project_code",
            "project_status",
            "region",
            "market_type",
            "reliability_grade",
            "source",
            "source_code",
        ],
        "common_base": "v5_context_terms",
        "placement": {
            "structured_filter": True,
            "search_text": False,
            "embedding_text": True,
        },
        "purpose": "Measure the incremental value of domain metadata in embedding text.",
    },
]

SOURCE_VALUE = "Vinhomes Market và nguồn công khai"
OWNER_ID = "eval-owner"
NOTEBOOK_ID = "eval-notebook"


def _clean(value: object) -> str:
    return " ".join(str(value or "").split())


def _normalize(value: object) -> str:
    text = unicodedata.normalize("NFD", _clean(value).casefold())
    text = "".join(char for char in text if unicodedata.category(char) != "Mn")
    text = re.sub(r"\bvinhomes\b", "", text)
    return re.sub(r"[^a-z0-9]+", " ", text).strip()


QUERY_STYLES = (
    "canonical",
    "conversational",
    "concise",
    "no_accents",
    "abbreviated",
    "light_typo",
)


def _without_accents(value: str) -> str:
    text = unicodedata.normalize("NFD", value)
    text = "".join(char for char in text if unicodedata.category(char) != "Mn")
    return text.replace("đ", "d").replace("Đ", "D")


def _style_query(query: str, index: int) -> tuple[str, str]:
    style = QUERY_STYLES[(index - 1) % len(QUERY_STYLES)]
    if style == "conversational":
        return f"Cho mình hỏi: {query[0].lower()}{query[1:]}", style
    if style == "concise":
        return f"Cần xác minh: {query}", style
    if style == "no_accents":
        return _without_accents(query), style
    if style == "abbreviated":
        replacements = {
            "tài liệu": "TL",
            "Tài liệu": "TL",
            "dự án": "DA",
            "Dự án": "DA",
            "Vinhomes": "VH",
        }
        for source, replacement in replacements.items():
            query = query.replace(source, replacement)
        return query, style
    if style == "light_typo":
        replacements = (
            ("tiện ích", "tien ich"),
            ("tài liệu", "tai lieu"),
            ("mức giá", "muc gia"),
        )
        for source, replacement in replacements:
            if source in query:
                return query.replace(source, replacement, 1), style
        words = query.split()
        for word_index, word in enumerate(words):
            misspelled = _without_accents(word)
            if misspelled != word:
                words[word_index] = misspelled
                return " ".join(words), style
    return query, style


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _write_json(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(value, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )


def _write_jsonl(path: Path, rows: Iterable[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="\n") as handle:
        for row in rows:
            handle.write(json.dumps(row, ensure_ascii=False, sort_keys=True) + "\n")


def _jsonl_sha256(rows: Iterable[dict[str, Any]]) -> str:
    digest = hashlib.sha256()
    for row in rows:
        line = json.dumps(row, ensure_ascii=False, sort_keys=True) + "\n"
        digest.update(line.encode("utf-8"))
    return digest.hexdigest()


def _json_sha256(value: Any) -> str:
    payload = json.dumps(value, ensure_ascii=False, indent=2) + "\n"
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()


def _write_csv(path: Path, rows: list[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if not rows:
        path.write_text("", encoding="utf-8")
        return
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0]))
        writer.writeheader()
        writer.writerows(rows)


def _table_rows(document: Document, table_index: int) -> list[list[str]]:
    return [[_clean(cell.text) for cell in row.cells] for row in document.tables[table_index].rows]


def _first_phrase(value: str) -> str:
    # Dots are thousands separators in values such as 5.000 m2 and 51.000 m2.
    phrase = value.split(";", maxsplit=1)[0].strip().rstrip(".")
    return phrase or value


def _extract_price_facts(source_dir: Path) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    facts: list[dict[str, Any]] = []
    summaries: list[dict[str, Any]] = []
    for year, filename in PRICE_DOCUMENTS.items():
        document = Document(source_dir / filename)
        for row_number, row in enumerate(_table_rows(document, 5)[1:], start=1):
            if len(row) < 6:
                continue
            facts.append(
                {
                    "fact_id": f"price_{year}_{row_number:02d}",
                    "filename": filename,
                    "year": year,
                    "project": row[0],
                    "product": row[1],
                    "price": row[2],
                    "time_market": row[3],
                    "grade": row[4],
                    "source_code": row[5],
                }
            )
        for row_number, row in enumerate(_table_rows(document, 2)[1:], start=1):
            if len(row) < 3:
                continue
            summaries.append(
                {
                    "fact_id": f"summary_{year}_{row_number:02d}",
                    "filename": filename,
                    "year": year,
                    "indicator": row[0],
                    "value": row[1],
                    "interpretation": row[2],
                }
            )
    return facts, summaries


def _extract_utility_facts(source_dir: Path) -> list[dict[str, Any]]:
    document = Document(source_dir / UTILITY_DOCUMENT)
    overview: dict[str, dict[str, str]] = {}
    for row in _table_rows(document, 2)[1:]:
        if len(row) < 6:
            continue
        overview[row[0]] = {
            "project": row[1],
            "region": row[2],
            "scale": row[3],
            "start": row[4],
            "status": row[5],
        }

    headings = [
        _clean(paragraph.text)
        for paragraph in document.paragraphs
        if paragraph._p.pPr is not None
        and paragraph._p.pPr.style == "Heading3"
        and "•" in paragraph.text
    ]
    detail_tables = document.tables[5 : 5 + len(headings)]
    if len(headings) != 18 or len(detail_tables) != 18:
        raise ValueError(
            f"Expected 18 utility project sections, got headings={len(headings)}, "
            f"tables={len(detail_tables)}"
        )

    facts: list[dict[str, Any]] = []
    for heading, table in zip(headings, detail_tables, strict=True):
        code, project = (_clean(value) for value in heading.split("•", maxsplit=1))
        row = [_clean(cell.text) for cell in table.rows[1].cells]
        base = overview.get(code)
        if base is None:
            raise ValueError(f"Missing utility overview row for {code}")
        facts.append(
            {
                "fact_id": f"utility_{code.lower()}",
                "filename": UTILITY_DOCUMENT,
                "heading": heading,
                "code": code,
                "project": project,
                "region": base["region"],
                "scale": base["scale"],
                "start": base["start"],
                "status": base["status"],
                "location_scale": row[0],
                "amenities": row[1],
                "caveat": row[2],
                "amenity_term": _first_phrase(row[1]),
            }
        )
    return facts


def _numeric_fact_integrity_audit(facts: list[dict[str, Any]]) -> list[dict[str, Any]]:
    rows = []
    for fact in facts:
        source_clause = str(fact["amenities"]).split(";", maxsplit=1)[0].strip().rstrip(".")
        extracted_anchor = str(fact["amenity_term"])
        numeric_phrases = re.findall(
            r"\b\d+(?:[.,]\d+)*(?:\s*(?:m²|ha|phòng|căn))?",
            source_clause,
            flags=re.IGNORECASE,
        )
        rows.append(
            {
                "fact_id": fact["fact_id"],
                "project_code": fact["code"],
                "project": fact["project"],
                "source_first_clause": source_clause,
                "extracted_anchor": extracted_anchor,
                "numeric_phrases": " | ".join(numeric_phrases),
                "contains_thousands_separator": bool(
                    re.search(r"\b\d{1,3}\.\d{3}\b", source_clause)
                ),
                "status": "match" if extracted_anchor == source_clause else "mismatch",
            }
        )
    return rows


def _extract_plan_facts(
    source_dir: Path,
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    document = Document(source_dir / PLAN_DOCUMENT)
    statuses: list[dict[str, Any]] = []
    for row_number, row in enumerate(_table_rows(document, 3)[1:], start=1):
        if len(row) < 6:
            continue
        statuses.append(
            {
                "fact_id": f"status_{row_number:02d}",
                "filename": PLAN_DOCUMENT,
                "project": row[0],
                "region": row[1],
                "scale": row[2],
                "start": row[3],
                "status": row[4],
                "source_code": row[5],
            }
        )
    milestones: list[dict[str, Any]] = []
    for row_number, row in enumerate(_table_rows(document, 5)[1:], start=1):
        if len(row) < 5:
            continue
        milestones.append(
            {
                "fact_id": f"milestone_{row_number:02d}",
                "filename": PLAN_DOCUMENT,
                "date": row[0],
                "project": row[1],
                "description": row[2],
                "milestone_type": row[3],
                "source_code": row[4],
            }
        )
    return statuses, milestones


def _price_profile(year: int, facts: list[dict[str, Any]]) -> dict[str, Any]:
    filename = PRICE_DOCUMENTS[year]
    projects = list(dict.fromkeys(fact["project"] for fact in facts if fact["year"] == year))
    grades = list(dict.fromkeys(fact["grade"] for fact in facts if fact["year"] == year))
    source_codes = list(
        dict.fromkeys(fact["source_code"] for fact in facts if fact["year"] == year)
    )
    status = "latest" if year == 2026 else "historical"
    current = {
        "title": filename,
        "document_type": "market_report",
        "source_kind": "docx",
        "language": "vi",
    }
    gold = {
        **current,
        "domain": "real_estate_market",
        "source": SOURCE_VALUE,
        "year": year,
        "data_period": str(year),
        "document_version": year - 2022,
        "effective_status": "current" if year == 2026 else "historical",
        "lifecycle_status": status,
        "as_of_date": "2026-08-04",
        "keyword_aliases": [f"Giá nhà Vinhomes {year}", "hồ sơ giá nhà", "bảng giá"],
    }
    return {
        "current_document_metadata": current,
        "document_metadata": gold,
        "rules": [
            {
                "section_titles": ["1. Tóm tắt điều hành"],
                "metadata": {
                    "clause_type": "executive_summary",
                    "contextual_summary": f"Tóm tắt các chỉ báo giá Vinhomes năm {year}.",
                    "contextual_search_terms": [f"giá nhà {year}", "chỉ báo giá"],
                },
            },
            {
                "section_titles": ["1. Tóm tắt điều hành"],
                "contains_any": ["| Chỉ báo | Mức ghi nhận | Cách đọc |"],
                "metadata": {"table_header": "Chỉ báo | Mức ghi nhận | Cách đọc"},
            },
            {
                "section_titles": ["2. Phạm vi và phương pháp"],
                "metadata": {
                    "clause_type": "methodology",
                    "contextual_summary": (
                        f"Phạm vi nguồn, độ tin cậy và cách diễn giải dữ liệu giá năm {year}."
                    ),
                    "contextual_search_terms": ["hạng A", "hạng B", "hạng C"],
                },
            },
            {
                "section_titles": ["2. Phạm vi và phương pháp"],
                "contains_any": ["| Hạng | Định nghĩa | Cách sử dụng |"],
                "metadata": {"table_header": "Hạng | Định nghĩa | Cách sử dụng"},
            },
            {
                "section_titles": ["3. Bảng dữ liệu giá chi tiết"],
                "contains_any": ["| Dự án | Sản phẩm / phân khu | Mức giá ghi nhận |"],
                "metadata": {
                    "clause_type": "price_observation_table",
                    "policy_field": "market_price",
                    "table_header": (
                        "Dự án | Sản phẩm / phân khu | Mức giá ghi nhận | "
                        "Thời điểm / loại | Hạng | Nguồn"
                    ),
                    "project_name": projects,
                    "market_type": ["primary", "survey", "transfer"],
                    "reliability_grade": grades,
                    "source_code": source_codes,
                    "contextual_summary": (
                        f"Bảng dữ liệu giá chi tiết theo dự án và sản phẩm trong hồ sơ năm {year}."
                    ),
                    "contextual_search_terms": [*projects, f"giá Vinhomes {year}"],
                },
            },
            {
                "section_titles": ["4. Nhận định và cách diễn giải"],
                "metadata": {
                    "clause_type": "interpretation",
                    "contextual_summary": f"Hướng dẫn diễn giải và so sánh dữ liệu giá năm {year}.",
                    "contextual_search_terms": ["giá all-in", "sơ cấp", "chuyển nhượng"],
                },
            },
            {
                "section_titles": ["5. Checklist xác minh trước giao dịch"],
                "metadata": {
                    "clause_type": "verification_checklist",
                    "contextual_summary": "Checklist xác minh giá, pháp lý và điều kiện giao dịch.",
                    "contextual_search_terms": ["VAT", "kinh phí bảo trì", "lịch thanh toán"],
                },
            },
        ],
    }


def _utility_profile(facts: list[dict[str, Any]]) -> dict[str, Any]:
    current = {
        "title": UTILITY_DOCUMENT,
        "document_type": "amenity_catalog",
        "source_kind": "docx",
        "language": "vi",
    }
    gold = {
        **current,
        "domain": "real_estate_projects",
        "source": SOURCE_VALUE,
        "year": 2026,
        "data_period": "2026",
        "document_version": 1,
        "effective_status": "current",
        "lifecycle_status": "latest",
        "as_of_date": "2026-08-04",
        "keyword_aliases": ["tiện ích Vinhomes", "bản đồ tiện ích", "18 dự án"],
    }
    rules: list[dict[str, Any]] = [
        {
            "section_titles": ["1. Toàn cảnh danh mục"],
            "contains_any": ["| Mã | Dự án | Địa bàn | Quy mô | Khởi công | Trạng thái |"],
            "metadata": {
                "clause_type": "project_catalog",
                "table_header": "Mã | Dự án | Địa bàn | Quy mô | Khởi công | Trạng thái",
                "project_name": [fact["project"] for fact in facts],
                "project_code": [fact["code"] for fact in facts],
                "region": list(dict.fromkeys(fact["region"] for fact in facts)),
                "project_status": list(dict.fromkeys(fact["status"] for fact in facts)),
                "contextual_summary": "Danh mục 18 dự án cùng địa bàn, quy mô và trạng thái.",
                "contextual_search_terms": ["P01", "P18", "trạng thái dự án"],
            },
        },
        {
            "section_titles": ["2. Khung đánh giá tiện ích"],
            "metadata": {
                "clause_type": "amenity_evaluation_framework",
                "contextual_summary": "Khung kiểm tra mức độ vận hành của các nhóm tiện ích.",
                "contextual_search_terms": ["giáo dục", "y tế", "thương mại", "giao thông"],
            },
        },
        {
            "section_titles": ["2. Khung đánh giá tiện ích"],
            "contains_any": ["| Nhóm | Nội dung cần kiểm tra | Chỉ báo thực tế |"],
            "metadata": {"table_header": "Nhóm | Nội dung cần kiểm tra | Chỉ báo thực tế"},
        },
        {
            "section_titles": ["4. Cụm dự án du lịch–nghỉ dưỡng"],
            "metadata": {
                "clause_type": "tourism_cluster_comparison",
                "contextual_summary": "So sánh hạt nhân trải nghiệm của cụm dự án nghỉ dưỡng.",
                "contextual_search_terms": ["Pearl Bay", "Hải Vân Bay", "Green Paradise"],
            },
        },
        {
            "section_titles": ["4. Cụm dự án du lịch–nghỉ dưỡng"],
            "contains_any": ["| Dự án | Hạt nhân trải nghiệm | Định vị sử dụng |"],
            "metadata": {
                "table_header": (
                    "Dự án | Hạt nhân trải nghiệm | Định vị sử dụng | Điểm phải xác minh"
                )
            },
        },
        {
            "section_titles": ["5. Checklist khảo sát tiện ích"],
            "metadata": {
                "clause_type": "amenity_checklist",
                "contextual_summary": "Checklist khảo sát tiện ích tại hiện trường.",
                "contextual_search_terms": ["đã vận hành", "đang xây", "quy hoạch"],
            },
        },
    ]
    for fact in facts:
        rules.append(
            {
                "section_titles": [fact["heading"]],
                "contains_any": ["| Địa bàn / quy mô | Tiện ích nổi bật |"],
                "metadata": {
                    "clause_type": "project_amenities",
                    "table_header": ("Địa bàn / quy mô | Tiện ích nổi bật | Tình trạng cần lưu ý"),
                    "project_name": fact["project"],
                    "project_code": fact["code"],
                    "region": fact["region"],
                    "project_status": fact["status"],
                    "contextual_summary": (
                        f"Tiện ích nổi bật và lưu ý triển khai của {fact['project']} "
                        f"tại {fact['region']}."
                    ),
                    "contextual_search_terms": [
                        fact["project"],
                        fact["code"],
                        fact["region"],
                        fact["amenity_term"],
                    ],
                },
            }
        )
    return {
        "current_document_metadata": current,
        "document_metadata": gold,
        "rules": rules,
    }


def _plan_profile(
    statuses: list[dict[str, Any]], milestones: list[dict[str, Any]]
) -> dict[str, Any]:
    current = {
        "title": PLAN_DOCUMENT,
        "document_type": "implementation_plan",
        "source_kind": "docx",
        "language": "vi",
    }
    gold = {
        **current,
        "domain": "real_estate_projects",
        "source": SOURCE_VALUE,
        "year": 2026,
        "data_period": "through 2026-08-04",
        "document_version": 1,
        "effective_status": "current",
        "lifecycle_status": "latest",
        "as_of_date": "2026-08-04",
        "keyword_aliases": ["kế hoạch xây dựng", "tiến độ Vinhomes", "mốc triển khai"],
    }
    cluster_terms = {
        "Vùng Thủ đô": ["Wonder City", "trung tâm triển lãm", "kết nối vùng"],
        "Hải Phòng–Quảng Ninh": ["Royal Island", "Global Gate Hạ Long", "hạ tầng vịnh"],
        "Duyên hải miền Trung": ["Pearl Bay", "Hải Vân Bay", "Q4/2027"],
        "TP.HCM và vùng phụ cận": ["Grand Park", "Green City", "Green Paradise"],
    }
    rules: list[dict[str, Any]] = [
        {
            "section_titles": ["1. Tóm tắt điều hành"],
            "metadata": {
                "clause_type": "executive_summary",
                "contextual_summary": "Tóm tắt chu kỳ mở rộng và các cụm triển khai 2024–2026.",
                "contextual_search_terms": ["2024–2026", "đại đô thị", "du lịch ven biển"],
            },
        },
        {
            "section_titles": ["1. Tóm tắt điều hành"],
            "contains_any": ["| Giai đoạn | Mốc chính | Ý nghĩa triển khai |"],
            "metadata": {"table_header": "Giai đoạn | Mốc chính | Ý nghĩa triển khai"},
        },
        {
            "section_titles": ["2. Danh mục và trạng thái công khai"],
            "contains_any": ["| Dự án | Địa bàn | Quy mô | Bắt đầu | Trạng thái | Nguồn |"],
            "metadata": {
                "clause_type": "project_status_catalog",
                "table_header": "Dự án | Địa bàn | Quy mô | Bắt đầu | Trạng thái | Nguồn",
                "project_name": [fact["project"] for fact in statuses],
                "project_status": list(dict.fromkeys(fact["status"] for fact in statuses)),
                "region": list(dict.fromkeys(fact["region"] for fact in statuses)),
                "source_code": list(dict.fromkeys(fact["source_code"] for fact in statuses)),
                "contextual_summary": "Danh mục 18 dự án và trạng thái công khai tại 04/08/2026.",
                "contextual_search_terms": ["đang xây dựng", "đã hoàn thành", "18 dự án"],
            },
        },
        {
            "section_titles": ["3. Các mốc triển khai trọng điểm"],
            "contains_any": ["| Thời điểm | Dự án | Nội dung | Loại mốc | Nguồn |"],
            "metadata": {
                "clause_type": "project_milestones",
                "table_header": "Thời điểm | Dự án | Nội dung | Loại mốc | Nguồn",
                "project_name": list(dict.fromkeys(fact["project"] for fact in milestones)),
                "source_code": list(dict.fromkeys(fact["source_code"] for fact in milestones)),
                "contextual_summary": (
                    "Các mốc khởi công, ra mắt và mục tiêu triển khai trọng điểm."
                ),
                "contextual_search_terms": [
                    "khởi công",
                    "ra mắt",
                    "mục tiêu hoàn thành",
                ],
            },
        },
        {
            "section_titles": ["5. Trọng tâm đô thị du lịch–nghỉ dưỡng"],
            "metadata": {
                "clause_type": "tourism_project_progress",
                "contextual_summary": "Trạng thái và mốc cần theo dõi của các dự án du lịch.",
                "contextual_search_terms": ["Pearl Bay", "Global Gate Hạ Long", "Hải Vân Bay"],
            },
        },
        {
            "section_titles": ["5. Trọng tâm đô thị du lịch–nghỉ dưỡng"],
            "contains_any": ["| Dự án | Hợp phần chủ đạo | Trạng thái công khai |"],
            "metadata": {
                "table_header": (
                    "Dự án | Hợp phần chủ đạo | Trạng thái công khai | Mốc cần theo dõi"
                )
            },
        },
        {
            "section_titles": ["6. Ma trận theo dõi tiến độ"],
            "metadata": {
                "clause_type": "progress_monitoring_matrix",
                "contextual_summary": "Ma trận tài liệu, tần suất và dấu hiệu cảnh báo tiến độ.",
                "contextual_search_terms": ["pháp lý", "thi công", "bàn giao", "vận hành"],
            },
        },
        {
            "section_titles": ["6. Ma trận theo dõi tiến độ"],
            "contains_any": ["| Chỉ tiêu | Tài liệu cần lấy | Tần suất | Dấu hiệu cảnh báo |"],
            "metadata": {
                "table_header": "Chỉ tiêu | Tài liệu cần lấy | Tần suất | Dấu hiệu cảnh báo"
            },
        },
    ]
    for section, terms in cluster_terms.items():
        rules.append(
            {
                "section_titles": [section],
                "metadata": {
                    "clause_type": "deployment_cluster_analysis",
                    "region": section,
                    "contextual_summary": f"Phân tích cụm triển khai {section}.",
                    "contextual_search_terms": terms,
                },
            }
        )
    return {
        "current_document_metadata": current,
        "document_metadata": gold,
        "rules": rules,
    }


def _build_gold_config(
    price_facts: list[dict[str, Any]],
    utility_facts: list[dict[str, Any]],
    statuses: list[dict[str, Any]],
    milestones: list[dict[str, Any]],
) -> dict[str, Any]:
    base = json.loads(BASE_GOLD_PATH.read_text(encoding="utf-8"))
    documents = dict(base["documents"])
    for year in PRICE_DOCUMENTS:
        documents[PRICE_DOCUMENTS[year]] = _price_profile(year, price_facts)
    documents[UTILITY_DOCUMENT] = _utility_profile(utility_facts)
    documents[PLAN_DOCUMENT] = _plan_profile(statuses, milestones)
    return {
        "schema_version": "3.0",
        "benchmark_kind": "real_document_metadata_benchmark",
        "annotation_policy": (
            "Evidence is extracted from uploaded source tables and anchored to deterministic "
            "chunk IDs. Human review remains required before treating the set as final gold."
        ),
        "documents": documents,
    }


def _find_chunk(
    chunks_by_document: dict[str, list[EvalChunk]],
    filename: str,
    *,
    section_title: str | None = None,
    content_kind: str | None = None,
    terms: Iterable[str] = (),
) -> EvalChunk:
    candidates = chunks_by_document.get(filename, [])
    if section_title is not None:
        candidates = [
            chunk
            for chunk in candidates
            if _normalize(chunk.current_metadata.get("section_title")) == _normalize(section_title)
        ]
    if content_kind is not None:
        candidates = [
            chunk
            for chunk in candidates
            if _normalize(chunk.current_metadata.get("content_kind")) == _normalize(content_kind)
        ]
    required = [_clean(term) for term in terms if _clean(term)]
    candidates = [
        chunk for chunk in candidates if all(_contains(chunk.text, term) for term in required)
    ]
    if len(candidates) != 1:
        raise ValueError(
            f"Expected one chunk for {filename}, section={section_title!r}, "
            f"kind={content_kind!r}, terms={required!r}; found {len(candidates)}"
        )
    return candidates[0]


def _attach_chunks(
    corpus: list[EvalChunk],
    price_facts: list[dict[str, Any]],
    summaries: list[dict[str, Any]],
    utility_facts: list[dict[str, Any]],
    statuses: list[dict[str, Any]],
    milestones: list[dict[str, Any]],
) -> dict[str, Any]:
    by_document: dict[str, list[EvalChunk]] = defaultdict(list)
    for chunk in corpus:
        by_document[chunk.document_title].append(chunk)

    for fact in price_facts:
        chunk = _find_chunk(
            by_document,
            fact["filename"],
            section_title="3. Bảng dữ liệu giá chi tiết",
            content_kind="table",
            terms=[fact["project"], fact["product"], fact["price"]],
        )
        fact["chunk_id"] = chunk.id
    for fact in summaries:
        chunk = _find_chunk(
            by_document,
            fact["filename"],
            section_title="1. Tóm tắt điều hành",
            content_kind="table",
            terms=[fact["indicator"], fact["value"]],
        )
        fact["chunk_id"] = chunk.id
    for fact in utility_facts:
        chunk = _find_chunk(
            by_document,
            UTILITY_DOCUMENT,
            section_title=fact["heading"],
            content_kind="table",
            terms=[fact["amenity_term"], fact["caveat"]],
        )
        fact["chunk_id"] = chunk.id
    for fact in statuses:
        chunk = _find_chunk(
            by_document,
            PLAN_DOCUMENT,
            section_title="2. Danh mục và trạng thái công khai",
            content_kind="table",
            terms=[fact["project"], fact["status"]],
        )
        fact["chunk_id"] = chunk.id
    for fact in milestones:
        chunk = _find_chunk(
            by_document,
            PLAN_DOCUMENT,
            section_title="3. Các mốc triển khai trọng điểm",
            content_kind="table",
            terms=[fact["date"], fact["project"], fact["description"]],
        )
        fact["chunk_id"] = chunk.id
    return {"by_document": by_document, "by_id": {chunk.id: chunk for chunk in corpus}}


def _condition(field: str, value: object, op: str = "eq") -> dict[str, Any]:
    return {"field": field, "op": op, "value": value}


def _case(
    *,
    primary_slice: str,
    index: int,
    query: str,
    chunks_by_id: dict[str, EvalChunk],
    relevant_chunk_ids: list[str] | None = None,
    relevant_chunk_groups: list[list[str]] | None = None,
    protected_chunk_ids: list[str] | None = None,
    target_type: str = "single",
    answerable: bool = True,
    difficulty: str = "medium",
    required_metadata_fields: list[str] | None = None,
    metadata_conditions: list[dict[str, Any]] | None = None,
    must_include_terms: list[str] | None = None,
    must_not_include_terms: list[str] | None = None,
    forbidden_chunk_ids: list[str] | None = None,
    forbidden_document_titles: list[str] | None = None,
    must_cite_document_titles: list[str] | None = None,
    must_not_cite_document_titles: list[str] | None = None,
    extra_slices: list[str] | None = None,
    query_context: dict[str, Any] | None = None,
    evidence_fact_ids: list[str] | None = None,
    unsupported_field_policy: str = "fail_closed",
    response_class: str | None = None,
    permission_pair_id: str | None = None,
    case_suffix: str | None = None,
    scenario_id: str | None = None,
    table_target: dict[str, Any] | None = None,
    conflict_target: dict[str, Any] | None = None,
) -> dict[str, Any]:
    suffix = f"_{case_suffix}" if case_suffix else ""
    case_id = f"real_{primary_slice}_{index:02d}{suffix}"
    styled_query, query_style = _style_query(query, index)
    relevant = list(dict.fromkeys(relevant_chunk_ids or []))
    protected = list(dict.fromkeys(protected_chunk_ids or []))
    groups = relevant_chunk_groups or ([[chunk_id] for chunk_id in relevant] if relevant else [])
    all_reference_ids = [*relevant, *protected]
    titles = list(
        dict.fromkeys(
            chunks_by_id[chunk_id].document_title
            for chunk_id in all_reference_ids
            if chunk_id in chunks_by_id
        )
    )
    slices = list(dict.fromkeys([primary_slice, *(extra_slices or [])]))
    relevant_titles = list(
        dict.fromkeys(
            chunks_by_id[chunk_id].document_title
            for chunk_id in relevant
            if chunk_id in chunks_by_id
        )
    )
    expected_title = relevant_titles[0] if len(relevant_titles) == 1 else ""
    required_citations = (
        list(dict.fromkeys(must_cite_document_titles))
        if must_cite_document_titles is not None
        else (relevant_titles if answerable else [])
    )
    forbidden_citations = (
        list(dict.fromkeys(must_not_cite_document_titles))
        if must_not_cite_document_titles is not None
        else ([] if answerable else titles)
    )
    expected = {
        "target_type": target_type,
        "document_title": expected_title,
        "page": (table_target or {}).get("page"),
        "page_tolerance": 0,
        "must_include_terms": must_include_terms or [],
        "must_not_include_terms": must_not_include_terms or [],
        "should_include_terms": [],
        "must_cite_document_titles": required_citations,
        "must_not_cite_document_titles": forbidden_citations,
        "forbidden_document_titles": forbidden_document_titles or [],
        "forbidden_chunk_ids": forbidden_chunk_ids or [],
        "protected_chunk_ids": protected,
        "response_class": response_class or ("ANSWER_WITH_CITATIONS" if answerable else ""),
        **(table_target or {}),
        **(conflict_target or {}),
    }
    return {
        "id": case_id,
        "query_id": case_id,
        "query": styled_query,
        "query_style": query_style,
        "query_type": primary_slice,
        "category": primary_slice,
        "primary_slice": primary_slice,
        "benchmark_slices": slices,
        "scenario_id": scenario_id or (evidence_fact_ids[0] if evidence_fact_ids else case_id),
        "split": "dev" if index <= 6 else "test",
        "difficulty": difficulty,
        "answerable": answerable,
        "target_type": target_type,
        "source_file": titles[0] if titles else "__none__",
        "source_kind": "uploaded_real_documents",
        "domain": "real_estate",
        "metadata_focus": required_metadata_fields or [],
        "required_metadata_fields": required_metadata_fields or [],
        "relevant_doc_ids": list(
            dict.fromkeys(
                chunks_by_id[chunk_id].document_id
                for chunk_id in relevant
                if chunk_id in chunks_by_id
            )
        ),
        "relevant_doc_titles": relevant_titles,
        "relevant_chunk_ids": relevant,
        "relevant_chunk_groups": groups,
        "protected_chunk_ids": protected,
        "forbidden_chunk_ids": forbidden_chunk_ids or [],
        "retrieval_filters": {
            "metadata_conditions": metadata_conditions or [],
            "unsupported_field_policy": unsupported_field_policy,
        },
        "query_context": query_context
        or {
            "owner_id": OWNER_ID,
            "notebook_id": NOTEBOOK_ID,
            "groups": ["benchmark-reader"],
        },
        "expected_metadata": {
            "required_fields": required_metadata_fields or [],
            "metadata_conditions": metadata_conditions or [],
        },
        "expected": expected,
        "evidence_fact_ids": evidence_fact_ids or [],
        "permission_pair_id": permission_pair_id,
        "structured_subtype": "table" if table_target else None,
        "annotation_status": "evidence_anchored_pending_human_review",
        "security_simulated": target_type
        in {"permission", "permission_allowed", "permission_denied"},
        "notes": (
            "Auto-extracted from uploaded documents; approve the review CSV before final gold use."
        ),
    }


def _round_robin(groups: dict[str, list[dict[str, Any]]]) -> list[dict[str, Any]]:
    output: list[dict[str, Any]] = []
    queues = {key: list(value) for key, value in sorted(groups.items())}
    while any(queues.values()):
        for key in queues:
            if queues[key]:
                output.append(queues[key].pop(0))
    return output


def _table_target(
    chunk: EvalChunk,
    *,
    row_key: str,
    column_key: str,
    expected_cell_value: str,
    header_path: list[str],
) -> dict[str, Any]:
    if not chunk.table_identity or not chunk.source_block_ids:
        raise ValueError(f"Chunk {chunk.id} has no atomic table provenance")
    return {
        "table_id": chunk.table_identity,
        "table_location": chunk.table_location,
        "row_key": row_key,
        "column_key": column_key,
        "expected_cell_value": expected_cell_value,
        "header_path": header_path,
        "page": chunk.page_number,
        "page_is_physical": False,
        "source_block_ids": list(chunk.source_block_ids),
        "bbox": list(chunk.bbox),
    }


def _other_price_documents(price_facts: list[dict[str, Any]], fact: dict[str, Any]) -> list[str]:
    project = _normalize(fact["project"])
    return list(
        dict.fromkeys(
            other["filename"]
            for other in price_facts
            if _normalize(other["project"]) == project and other["filename"] != fact["filename"]
        )
    )


def _plan_evidence_items(
    by_document: dict[str, list[EvalChunk]],
) -> list[dict[str, Any]]:
    specs = [
        {
            "label": "document_identity",
            "section": "DOCX",
            "kind": "paragraph",
            "terms": ["Kế hoạch xây dựng và triển khai dự án Vinhomes", "04/08/2026"],
            "query": "Tài liệu kế hoạch bao quát các cụm trọng điểm đến ngày nào?",
        },
        {
            "label": "scope",
            "section": "DOCX",
            "kind": "paragraph",
            "terms": ["Tổng hợp trạng thái công khai", "18 dự án"],
            "query": "Phạm vi tài liệu tiến độ tổng hợp trạng thái công khai của bao nhiêu dự án?",
        },
        {
            "label": "as_of_note",
            "section": "DOCX",
            "kind": "table",
            "terms": ["MỐC DỮ LIỆU", "Giá và kế hoạch có thể thay đổi"],
            "query": (
                "Mốc dữ liệu của hồ sơ là ngày nào và tài liệu cảnh báo điều gì có thể thay đổi?"
            ),
        },
        {
            "label": "executive_summary",
            "section": "1. Tóm tắt điều hành",
            "kind": "paragraph",
            "terms": ["Chu kỳ 2024–2026", "mở rộng đồng thời ở ba hướng"],
            "query": "Chu kỳ 2024–2026 cho thấy Vinhomes mở rộng đồng thời theo những hướng nào?",
        },
        {
            "label": "phase_table",
            "section": "1. Tóm tắt điều hành",
            "kind": "table",
            "terms": ["2018–2022", "Hình thành nền tảng đại đô thị"],
            "query": "Giai đoạn 2018–2022 có ý nghĩa triển khai gì trong bảng tóm tắt?",
        },
        {
            "label": "status_table",
            "section": "2. Danh mục và trạng thái công khai",
            "kind": "table",
            "terms": ["Vinhomes Sài Gòn Park", "Đang xây dựng"],
            "query": "Bảng danh mục ghi Vinhomes Sài Gòn Park ở địa bàn và trạng thái nào?",
            "table_annotation": {
                "row_key": "Vinhomes Sài Gòn Park",
                "column_key": "Trạng thái",
                "expected_cell_value": "Đang xây dựng",
                "header_path": ["Dự án", "Trạng thái"],
            },
        },
        {
            "label": "status_limit",
            "section": "2. Danh mục và trạng thái công khai",
            "kind": "table",
            "terms": ["Giới hạn dữ liệu", "phần trăm hoàn thành"],
            "query": (
                "Giới hạn dữ liệu giải thích vì sao trạng thái đang xây dựng "
                "không cho biết điều gì?"
            ),
        },
        {
            "label": "milestone_table",
            "section": "3. Các mốc triển khai trọng điểm",
            "kind": "table",
            "terms": ["26/03/2025", "Green City"],
            "query": "Bảng mốc trọng điểm ghi sự kiện nào của Green City ngày 26/03/2025?",
            "table_annotation": {
                "row_key": "26/03/2025 | Green City",
                "column_key": "Nội dung",
                "expected_cell_value": "Khởi công dự án 197,2 ha tại Long An",
                "header_path": ["Thời điểm", "Dự án", "Nội dung"],
            },
        },
        {
            "label": "cluster_capital",
            "section": "Vùng Thủ đô",
            "kind": "paragraph",
            "terms": ["Vùng Thủ đô", "Wonder City"],
            "query": "Cụm Vùng Thủ đô gồm những dự án nào và ưu tiên triển khai gì?",
            "answer_terms": ["Ưu tiên hạ tầng đại đô thị", "kết nối vùng"],
        },
        {
            "label": "cluster_northeast",
            "section": "Hải Phòng–Quảng Ninh",
            "kind": "paragraph",
            "terms": ["Hải Phòng–Quảng Ninh", "Global Gate Hạ Long"],
            "query": "Cụm Hải Phòng–Quảng Ninh tập trung vào các dự án và trọng tâm nào?",
            "answer_terms": ["Kết hợp nhà ở thấp tầng", "tiến độ hạ tầng lớn"],
        },
        {
            "label": "cluster_central",
            "section": "Duyên hải miền Trung",
            "kind": "paragraph",
            "terms": ["Duyên hải miền Trung", "Hải Vân Bay"],
            "query": "Cụm Duyên hải miền Trung nêu các dự án và mốc cần theo dõi nào?",
            "answer_terms": ["Royal Park đã hoàn thành", "Q4/2027"],
        },
        {
            "label": "cluster_south",
            "section": "TP.HCM và vùng phụ cận",
            "kind": "paragraph",
            "terms": ["TP.HCM và vùng phụ cận", "Green Paradise"],
            "query": "Cụm TP.HCM và vùng phụ cận gồm các dự án nào và có yêu cầu theo dõi gì?",
            "answer_terms": [
                "Grand Park tiếp tục hoàn thiện theo phân khu",
                "hạ tầng giao thông",
            ],
        },
        {
            "label": "tourism_table",
            "section": "5. Trọng tâm đô thị du lịch–nghỉ dưỡng",
            "kind": "table",
            "terms": ["Pearl Bay", "Hợp phần chủ đạo"],
            "query": (
                "Bảng trọng tâm du lịch nêu hợp phần chủ đạo và trạng thái của Pearl Bay ra sao?"
            ),
            "table_annotation": {
                "row_key": "Pearl Bay",
                "column_key": "Hợp phần chủ đạo",
                "expected_cell_value": ("Bãi biển, cáp treo, VinWonders, Vinpearl Harbour, golf"),
                "header_path": ["Dự án", "Hợp phần chủ đạo"],
            },
        },
        {
            "label": "monitoring_matrix",
            "section": "6. Ma trận theo dõi tiến độ",
            "kind": "table",
            "terms": ["Pháp lý", "Dấu hiệu cảnh báo"],
            "query": (
                "Ma trận tiến độ yêu cầu lấy tài liệu pháp lý nào và theo dõi dấu hiệu cảnh báo gì?"
            ),
            "table_annotation": {
                "row_key": "Pháp lý",
                "column_key": "Tài liệu cần lấy",
                "expected_cell_value": (
                    "Quy hoạch, giao đất/thuê đất, giấy phép xây dựng, bảo lãnh"
                ),
                "header_path": ["Chỉ tiêu", "Tài liệu cần lấy"],
            },
        },
        {
            "label": "usage_principle",
            "section": "6. Ma trận theo dõi tiến độ",
            "kind": "table",
            "terms": ["Không suy ra ngày bàn giao", "dự kiến/mục tiêu"],
            "query": "Nguyên tắc sử dụng cấm suy ra ngày bàn giao trong trường hợp nào?",
        },
    ]
    output: list[dict[str, Any]] = []
    for index, spec in enumerate(specs, start=1):
        chunk = _find_chunk(
            by_document,
            PLAN_DOCUMENT,
            section_title=str(spec["section"]),
            content_kind=str(spec["kind"]),
            terms=spec["terms"],
        )
        output.append(
            {
                **spec,
                "fact_id": f"plan_evidence_{index:02d}_{spec['label']}",
                "filename": PLAN_DOCUMENT,
                "chunk_id": chunk.id,
            }
        )
    return output


def _build_testset(
    *,
    chunks: dict[str, Any],
    price_facts: list[dict[str, Any]],
    summaries: list[dict[str, Any]],
    utility_facts: list[dict[str, Any]],
    statuses: list[dict[str, Any]],
    milestones: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    by_id: dict[str, EvalChunk] = chunks["by_id"]
    by_document: dict[str, list[EvalChunk]] = chunks["by_document"]
    tests: list[dict[str, Any]] = []
    plan_items = _plan_evidence_items(by_document)
    plan_by_label = {item["label"]: item for item in plan_items}

    pair_counts = Counter(
        (_normalize(fact["project"]), _normalize(fact["product"])) for fact in price_facts
    )
    unique_price_facts = [
        fact
        for fact in price_facts
        if pair_counts[(_normalize(fact["project"]), _normalize(fact["product"]))] == 1
    ]
    if len(unique_price_facts) < 30:
        raise ValueError(f"Need 30 unambiguous price facts, got {len(unique_price_facts)}")
    first_price_by_year = {
        year: next(fact for fact in unique_price_facts if fact["year"] == year)
        for year in PRICE_DOCUMENTS
    }
    first_summary_by_year = {
        year: next(fact for fact in summaries if fact["year"] == year) for year in PRICE_DOCUMENTS
    }
    cluster_items = [item for item in plan_items if str(item["label"]).startswith("cluster_")]

    content_index = 1
    for fact in utility_facts:
        tests.append(
            _case(
                primary_slice="content_only",
                index=content_index,
                query=(
                    f"Ở {fact['region']}, dự án quy mô {fact['scale']} này có điểm nổi bật "
                    "nào và thông tin nào còn cần xác minh?"
                ),
                chunks_by_id=by_id,
                relevant_chunk_ids=[fact["chunk_id"]],
                difficulty="easy",
                must_include_terms=[fact["amenities"], fact["caveat"]],
                evidence_fact_ids=[fact["fact_id"]],
            )
        )
        content_index += 1
    for year, fact in first_price_by_year.items():
        tests.append(
            _case(
                primary_slice="content_only",
                index=content_index,
                query=(
                    f"{fact['project']}, sản phẩm {fact['product']} có mức giá ghi nhận bao nhiêu?"
                ),
                chunks_by_id=by_id,
                relevant_chunk_ids=[fact["chunk_id"]],
                difficulty="easy",
                must_include_terms=[fact["price"]],
                evidence_fact_ids=[fact["fact_id"], f"year_{year}"],
            )
        )
        content_index += 1
    for fact in first_summary_by_year.values():
        tests.append(
            _case(
                primary_slice="content_only",
                index=content_index,
                query=f"Chỉ báo '{fact['indicator']}' có mức ghi nhận và cách đọc như thế nào?",
                chunks_by_id=by_id,
                relevant_chunk_ids=[fact["chunk_id"]],
                difficulty="easy",
                must_include_terms=[fact["value"], fact["interpretation"]],
                evidence_fact_ids=[fact["fact_id"]],
            )
        )
        content_index += 1
    for item in cluster_items:
        tests.append(
            _case(
                primary_slice="content_only",
                index=content_index,
                query=str(item["query"]),
                chunks_by_id=by_id,
                relevant_chunk_ids=[item["chunk_id"]],
                difficulty="easy",
                must_include_terms=list(item.get("answer_terms", item["terms"])),
                evidence_fact_ids=[item["fact_id"]],
            )
        )
        content_index += 1
    if content_index != 31:
        raise ValueError(f"Expected 30 content-only cases, built {content_index - 1}")

    for index, fact in enumerate(utility_facts, start=1):
        tests.append(
            _case(
                primary_slice="explicit_filter",
                index=index,
                query=(
                    "Trong amenity catalog năm 2026, nguồn Vinhomes Market và nguồn công khai, "
                    f"mục {fact['code']} của {fact['project']} nêu tiện ích gì?"
                ),
                chunks_by_id=by_id,
                relevant_chunk_ids=[fact["chunk_id"]],
                difficulty="hard",
                required_metadata_fields=["year", "document_type", "source", "project_name"],
                metadata_conditions=[
                    _condition("year", 2026),
                    _condition("document_type", "amenity_catalog"),
                    _condition("source", SOURCE_VALUE),
                    _condition("project_name", fact["project"]),
                ],
                must_include_terms=[fact["amenity_term"], fact["caveat"]],
                forbidden_document_titles=[PLAN_DOCUMENT],
                evidence_fact_ids=[fact["fact_id"]],
            )
        )
    explicit_price_facts = _round_robin(
        {
            str(year): [fact for fact in price_facts if fact["year"] == year]
            for year in PRICE_DOCUMENTS
        }
    )[:12]
    for index, fact in enumerate(explicit_price_facts, start=19):
        tests.append(
            _case(
                primary_slice="explicit_filter",
                index=index,
                query=(
                    f"Trong hồ sơ market report năm {fact['year']}, nguồn Vinhomes Market và "
                    f"nguồn công khai, {fact['project']} - {fact['product']} có giá bao nhiêu?"
                ),
                chunks_by_id=by_id,
                relevant_chunk_ids=[fact["chunk_id"]],
                difficulty="hard",
                required_metadata_fields=["year", "document_type", "source", "project_name"],
                metadata_conditions=[
                    _condition("year", fact["year"]),
                    _condition("document_type", "market_report"),
                    _condition("source", SOURCE_VALUE),
                    _condition("project_name", fact["project"]),
                ],
                must_include_terms=[fact["project"], fact["product"], fact["price"]],
                forbidden_document_titles=_other_price_documents(price_facts, fact),
                evidence_fact_ids=[fact["fact_id"]],
            )
        )

    latest_price = [fact for fact in price_facts if fact["year"] == 2026]
    implicit_plan_labels = [
        "executive_summary",
        "phase_table",
        "status_table",
        "status_limit",
        "milestone_table",
        "cluster_capital",
        "cluster_northeast",
        "cluster_central",
        "cluster_south",
        "tourism_table",
    ]
    implicit_items: list[tuple[str, dict[str, Any]]] = (
        [("utility", fact) for fact in utility_facts]
        + [("plan", plan_by_label[label]) for label in implicit_plan_labels]
        + [("price", fact) for fact in latest_price[:2]]
    )
    if len(implicit_items) != 30:
        raise ValueError(f"Expected 30 implicit-filter cases, got {len(implicit_items)}")
    for index, (kind, fact) in enumerate(implicit_items, start=1):
        if kind == "price":
            query = (
                f"Theo hồ sơ giá mới nhất hiện có, {fact['project']} - {fact['product']} "
                "được ghi nhận ở mức nào?"
            )
            document_type = "market_report"
            terms = [fact["project"], fact["product"], fact["price"]]
        elif kind == "utility":
            query = f"Trong danh mục tiện ích hiện hành, {fact['project']} có tiện ích nào nổi bật?"
            document_type = "amenity_catalog"
            terms = [fact["amenity_term"], fact["caveat"]]
        else:
            query = f"Theo kế hoạch hiện hành mới nhất, {str(fact['query']).lower()}"
            document_type = "implementation_plan"
            terms = list(fact["terms"])
        tests.append(
            _case(
                primary_slice="implicit_filter",
                index=index,
                query=query,
                chunks_by_id=by_id,
                relevant_chunk_ids=[fact["chunk_id"]],
                difficulty="hard",
                required_metadata_fields=[
                    "lifecycle_status",
                    "effective_status",
                    "document_type",
                    "as_of_date",
                ],
                metadata_conditions=[
                    _condition("lifecycle_status", "latest"),
                    _condition("document_type", document_type),
                ],
                must_include_terms=terms,
                evidence_fact_ids=[fact["fact_id"]],
            )
        )

    status_chunk_id = plan_by_label["status_table"]["chunk_id"]
    for index, fact in enumerate(utility_facts[:15], start=1):
        tests.append(
            _case(
                primary_slice="cross_document_confusion",
                index=index,
                query=(
                    f"{fact['project']} có những tiện ích nổi bật nào và cần lưu ý tình trạng gì?"
                ),
                chunks_by_id=by_id,
                relevant_chunk_ids=[fact["chunk_id"]],
                difficulty="hard",
                required_metadata_fields=["document_type", "section_title", "project_name"],
                must_include_terms=[fact["amenity_term"], fact["caveat"]],
                forbidden_chunk_ids=[status_chunk_id],
                forbidden_document_titles=[PLAN_DOCUMENT],
                evidence_fact_ids=[fact["fact_id"]],
            )
        )
    for offset, fact in enumerate(plan_items, start=16):
        tests.append(
            _case(
                primary_slice="cross_document_confusion",
                index=offset,
                query=f"Theo tài liệu kế hoạch, {str(fact['query']).lower()}",
                chunks_by_id=by_id,
                relevant_chunk_ids=[fact["chunk_id"]],
                difficulty="hard",
                required_metadata_fields=["document_type", "section_title", "project_name"],
                must_include_terms=list(fact["terms"]),
                forbidden_document_titles=[UTILITY_DOCUMENT],
                evidence_fact_ids=[fact["fact_id"]],
            )
        )

    by_project: dict[str, dict[int, list[dict[str, Any]]]] = defaultdict(lambda: defaultdict(list))
    for fact in price_facts:
        by_project[_normalize(fact["project"])][fact["year"]].append(fact)
    version_groups: dict[str, list[dict[str, Any]]] = {}
    latest_groups: dict[str, list[dict[str, Any]]] = {}
    for project_key, by_year in by_project.items():
        years = sorted(by_year)
        pairs: list[dict[str, Any]] = []
        for old_index, old_year in enumerate(years):
            for new_year in years[old_index + 1 :]:
                for old_fact in by_year[old_year]:
                    for new_fact in by_year[new_year]:
                        pairs.append({"old": old_fact, "new": new_fact})
        if pairs:
            version_groups[project_key] = pairs
        if 2026 in by_year and any(year < 2026 for year in years):
            latest_groups[project_key] = [
                {"old": old_fact, "new": new_fact}
                for old_year in years
                if old_year < 2026
                for old_fact in by_year[old_year]
                for new_fact in by_year[2026]
            ]

    latest_pairs = _round_robin(latest_groups)[:10]
    if len(latest_pairs) < 10:
        raise ValueError(f"Need 10 latest-resolution pairs, got {len(latest_pairs)}")
    for index, pair in enumerate(latest_pairs, start=1):
        old = pair["old"]
        new = pair["new"]
        tests.append(
            _case(
                primary_slice="version_conflict",
                index=index,
                query=(
                    f"Hồ sơ hiện hành mới nhất của {new['project']} ghi {new['product']} "
                    f"ở mức nào? Không dùng số liệu bản {old['year']}."
                ),
                chunks_by_id=by_id,
                relevant_chunk_ids=[new["chunk_id"]],
                difficulty="hard",
                required_metadata_fields=[
                    "document_version",
                    "lifecycle_status",
                    "project_name",
                ],
                metadata_conditions=[
                    _condition("lifecycle_status", "latest"),
                    _condition("project_name", new["project"]),
                ],
                must_include_terms=[new["price"]],
                forbidden_chunk_ids=[old["chunk_id"]],
                forbidden_document_titles=[old["filename"]],
                must_not_cite_document_titles=[old["filename"]],
                response_class="LATEST_EFFECTIVE_VALUE",
                conflict_target={
                    "conflict_type": "latest_effective_resolution",
                    "resolution_rule": "prefer_latest_effective_document",
                    "candidate_claims": [
                        {
                            "document_title": old["filename"],
                            "year": old["year"],
                            "scope": old["product"],
                            "value": old["price"],
                        },
                        {
                            "document_title": new["filename"],
                            "year": new["year"],
                            "scope": new["product"],
                            "value": new["price"],
                        },
                    ],
                },
                evidence_fact_ids=[old["fact_id"], new["fact_id"]],
            )
        )

    temporal_pairs = _round_robin(version_groups)[:10]
    for index, pair in enumerate(temporal_pairs, start=11):
        old = pair["old"]
        new = pair["new"]
        tests.append(
            _case(
                primary_slice="version_conflict",
                index=index,
                query=(
                    f"Hai hồ sơ của {old['project']} ghi các mức khác nhau. Hãy đối chiếu "
                    f"{old['year']} và {new['year']}, giữ rõ sản phẩm và thời điểm của từng mức."
                ),
                chunks_by_id=by_id,
                relevant_chunk_ids=[old["chunk_id"], new["chunk_id"]],
                relevant_chunk_groups=[[old["chunk_id"]], [new["chunk_id"]]],
                target_type="multi_hop",
                difficulty="hard",
                required_metadata_fields=["year", "document_version", "project_name"],
                must_include_terms=[old["price"], new["price"]],
                extra_slices=["multi_hop"],
                response_class="CONFLICT_REQUIRES_TEMPORAL_QUALIFICATION",
                conflict_target={
                    "conflict_type": "temporal_claim_difference",
                    "resolution_rule": "report_each_claim_with_year_and_scope",
                    "candidate_claims": [
                        {
                            "document_title": old["filename"],
                            "year": old["year"],
                            "scope": old["product"],
                            "value": old["price"],
                        },
                        {
                            "document_title": new["filename"],
                            "year": new["year"],
                            "scope": new["product"],
                            "value": new["price"],
                        },
                    ],
                },
                evidence_fact_ids=[old["fact_id"], new["fact_id"]],
            )
        )

    scope_groups: dict[str, list[dict[str, Any]]] = {}
    for project_key, by_year in by_project.items():
        for year, facts in by_year.items():
            pairs = [
                {"left": left, "right": right}
                for left_index, left in enumerate(facts)
                for right in facts[left_index + 1 :]
                if left["price"] != right["price"]
            ]
            if pairs:
                scope_groups[f"{project_key}:{year}"] = pairs
    scope_pairs = _round_robin(scope_groups)[:10]
    if len(scope_pairs) < 10:
        raise ValueError(f"Need 10 source/scope conflict pairs, got {len(scope_pairs)}")
    for index, pair in enumerate(scope_pairs, start=21):
        left = pair["left"]
        right = pair["right"]
        tests.append(
            _case(
                primary_slice="version_conflict",
                index=index,
                query=(
                    f"Cùng hồ sơ {left['year']} của {left['project']}, hai dòng "
                    f"{left['product']} và {right['product']} ghi mức nào; vì sao không nên "
                    "gộp thành một giá duy nhất?"
                ),
                chunks_by_id=by_id,
                relevant_chunk_ids=[left["chunk_id"]],
                difficulty="hard",
                required_metadata_fields=["year", "project_name", "source_code"],
                metadata_conditions=[
                    _condition("year", left["year"]),
                    _condition("project_name", left["project"]),
                ],
                must_include_terms=[left["price"], right["price"]],
                response_class="CONFLICT_REQUIRES_SCOPE_QUALIFICATION",
                conflict_target={
                    "conflict_type": "same_period_scope_difference",
                    "resolution_rule": "preserve_product_and_source_scope",
                    "resolution_evaluation_status": "requires_answer_evaluator",
                    "candidate_claims": [
                        {
                            "source_code": left["source_code"],
                            "scope": left["product"],
                            "value": left["price"],
                        },
                        {
                            "source_code": right["source_code"],
                            "scope": right["product"],
                            "value": right["price"],
                        },
                    ],
                },
                evidence_fact_ids=[left["fact_id"], right["fact_id"]],
            )
        )

    for index, fact in enumerate(utility_facts, start=1):
        tests.append(
            _case(
                primary_slice="section_localization",
                index=index,
                query=(
                    f"Trong mục '{fact['heading']}', tiện ích nổi bật và lưu ý "
                    "được nêu như thế nào?"
                ),
                chunks_by_id=by_id,
                relevant_chunk_ids=[fact["chunk_id"]],
                required_metadata_fields=["section_title", "section_path", "project_code"],
                must_include_terms=[fact["amenity_term"], fact["caveat"]],
                evidence_fact_ids=[fact["fact_id"]],
            )
        )
    cluster_cases = [
        ("Vùng Thủ đô", ["Wonder City", "kết nối vùng"]),
        ("Hải Phòng–Quảng Ninh", ["Royal Island", "Global Gate Hạ Long"]),
        ("Duyên hải miền Trung", ["Hải Vân Bay", "Q4/2027"]),
        ("TP.HCM và vùng phụ cận", ["Grand Park", "Green Paradise"]),
    ]
    for offset, (section, terms) in enumerate(cluster_cases, start=19):
        chunk = _find_chunk(
            by_document,
            PLAN_DOCUMENT,
            section_title=section,
            content_kind="paragraph",
            terms=terms,
        )
        tests.append(
            _case(
                primary_slice="section_localization",
                index=offset,
                query=f"Trong mục '{section}', tài liệu ưu tiên và cảnh báo triển khai điều gì?",
                chunks_by_id=by_id,
                relevant_chunk_ids=[chunk.id],
                required_metadata_fields=["section_title", "section_path", "region"],
                must_include_terms=terms,
                evidence_fact_ids=[f"cluster_{offset - 18:02d}"],
            )
        )
    first_summary_by_year = {
        year: next(fact for fact in summaries if fact["year"] == year) for year in PRICE_DOCUMENTS
    }
    section_index = 23
    for year in PRICE_DOCUMENTS:
        fact = first_summary_by_year[year]
        tests.append(
            _case(
                primary_slice="section_localization",
                index=section_index,
                query=(
                    f"Trong mục '1. Tóm tắt điều hành' của hồ sơ giá {year}, "
                    f"chỉ báo '{fact['indicator']}' có mức nào?"
                ),
                chunks_by_id=by_id,
                relevant_chunk_ids=[fact["chunk_id"]],
                required_metadata_fields=["section_title", "section_path", "year"],
                must_include_terms=[fact["indicator"], fact["value"]],
                evidence_fact_ids=[fact["fact_id"]],
            )
        )
        section_index += 1
        checklist_chunk = _find_chunk(
            by_document,
            PRICE_DOCUMENTS[year],
            section_title="5. Checklist xác minh trước giao dịch",
            content_kind="paragraph",
            terms=["Xin bảng tính đúng mã căn", "lịch thanh toán"],
        )
        tests.append(
            _case(
                primary_slice="section_localization",
                index=section_index,
                query=(
                    f"Trong mục '5. Checklist xác minh trước giao dịch' của hồ sơ {year}, "
                    "người mua cần xin bảng tính và kiểm tra lịch gì?"
                ),
                chunks_by_id=by_id,
                relevant_chunk_ids=[checklist_chunk.id],
                required_metadata_fields=["section_title", "section_path", "year"],
                must_include_terms=["Xin bảng tính đúng mã căn", "lịch thanh toán"],
                evidence_fact_ids=[f"checklist_{year}"],
            )
        )
        section_index += 1

    table_plan_items = [
        plan_by_label[label]
        for label in ("status_table", "milestone_table", "tourism_table", "monitoring_matrix")
    ]
    table_items: list[tuple[str, dict[str, Any]]] = (
        [("utility", fact) for fact in utility_facts]
        + [("price", fact) for fact in first_price_by_year.values()]
        + [("summary", fact) for fact in first_summary_by_year.values()]
        + [("plan", fact) for fact in table_plan_items]
    )
    table_cases: list[dict[str, Any]] = []
    for index, (kind, fact) in enumerate(table_items, start=1):
        chunk = by_id[fact["chunk_id"]]
        if kind == "price":
            query = (
                f"Đọc dòng bảng của {fact['project']} - {fact['product']}: "
                "mức giá và hạng nguồn là gì?"
            )
            terms = [fact["project"], fact["price"], fact["grade"]]
            target = _table_target(
                chunk,
                row_key=f"{fact['project']} | {fact['product']}",
                column_key="Mức giá ghi nhận",
                expected_cell_value=fact["price"],
                header_path=["Dự án", "Sản phẩm / phân khu", "Mức giá ghi nhận"],
            )
        elif kind == "utility":
            query = f"Đọc bảng tiện ích của {fact['project']}: tiện ích chính và lưu ý là gì?"
            terms = [fact["amenities"], fact["caveat"]]
            target = _table_target(
                chunk,
                row_key=fact["project"],
                column_key="Tiện ích nổi bật",
                expected_cell_value=fact["amenities"],
                header_path=["Địa bàn / quy mô", "Tiện ích nổi bật"],
            )
        elif kind == "summary":
            query = (
                f"Đọc bảng tóm tắt năm {fact['year']}: chỉ báo {fact['indicator']} "
                "có mức ghi nhận nào?"
            )
            terms = [fact["indicator"], fact["value"]]
            target = _table_target(
                chunk,
                row_key=fact["indicator"],
                column_key="Mức ghi nhận",
                expected_cell_value=fact["value"],
                header_path=["Chỉ báo", "Mức ghi nhận"],
            )
        else:
            query = f"Đọc đúng bảng trong tài liệu kế hoạch: {str(fact['query']).lower()}"
            terms = list(fact["terms"])
            annotation = fact["table_annotation"]
            target = _table_target(chunk, **annotation)
        table_case = _case(
            primary_slice="table_structured",
            index=index,
            query=query,
            chunks_by_id=by_id,
            relevant_chunk_ids=[fact["chunk_id"]],
            difficulty="hard",
            required_metadata_fields=["content_kind", "table_header", "section_path"],
            must_include_terms=terms,
            evidence_fact_ids=[fact["fact_id"]],
            table_target=target,
        )
        table_cases.append(table_case)
        tests.append(table_case)

    statuses_by_project = {_normalize(fact["project"]): fact for fact in statuses}
    cluster_label_by_code = {
        "P01": "cluster_south",
        "P02": "cluster_northeast",
        "P03": "cluster_central",
        "P04": "cluster_northeast",
        "P05": "cluster_central",
        "P06": "cluster_south",
        "P07": "cluster_northeast",
        "P08": "cluster_south",
        "P09": "cluster_capital",
        "P10": "cluster_central",
        "P11": "cluster_capital",
        "P12": "status_table",
        "P13": "cluster_northeast",
        "P14": "cluster_capital",
        "P15": "cluster_capital",
        "P16": "cluster_capital",
        "P17": "cluster_south",
        "P18": "cluster_capital",
    }

    def plan_item_for_utility(utility: dict[str, Any]) -> tuple[dict[str, Any], list[str]]:
        item = plan_by_label[cluster_label_by_code[utility["code"]]]
        if utility["code"] != "P12":
            return item, list(item["terms"])
        status = statuses_by_project[_normalize(utility["project"])]
        return item, [status["project"], status["start"], status["status"]]

    for index, utility in enumerate(utility_facts, start=1):
        if _normalize(utility["project"]) not in statuses_by_project:
            raise ValueError(f"No plan status for utility project {utility['project']}")
        plan_item, plan_terms = plan_item_for_utility(utility)
        tests.append(
            _case(
                primary_slice="multi_hop",
                index=index,
                query=(
                    f"Kết hợp hồ sơ tiện ích và phân tích triển khai: {utility['project']} có "
                    "tiện ích nổi bật nào và thuộc cụm hoặc trạng thái kế hoạch nào?"
                ),
                chunks_by_id=by_id,
                relevant_chunk_ids=[utility["chunk_id"], plan_item["chunk_id"]],
                relevant_chunk_groups=[[utility["chunk_id"]], [plan_item["chunk_id"]]],
                target_type="multi_hop",
                difficulty="hard",
                required_metadata_fields=[
                    "document_type",
                    "project_name",
                    "project_status",
                    "section_title",
                ],
                must_include_terms=[utility["amenity_term"], *plan_terms],
                evidence_fact_ids=[utility["fact_id"], plan_item["fact_id"]],
            )
        )
    utility_by_project = {_normalize(fact["project"]): fact for fact in utility_facts}
    triple_facts: list[tuple[dict[str, Any], dict[str, Any], dict[str, Any]]] = []
    seen_projects: set[str] = set()
    for price in latest_price:
        key = _normalize(price["project"])
        if key in seen_projects or key not in utility_by_project:
            continue
        seen_projects.add(key)
        utility = utility_by_project[key]
        plan_item, _ = plan_item_for_utility(utility)
        triple_facts.append((price, utility, plan_item))
    if len(triple_facts) < 12:
        raise ValueError(f"Need 12 price+utility+plan multi-hop projects, got {len(triple_facts)}")
    for index, (price, utility, plan_item) in enumerate(triple_facts[:12], start=19):
        _, plan_terms = plan_item_for_utility(utility)
        tests.append(
            _case(
                primary_slice="multi_hop",
                index=index,
                query=(
                    f"Tổng hợp {price['project']}: giá {price['product']} năm 2026, một tiện ích "
                    "nổi bật và cụm triển khai hoặc trọng tâm kế hoạch là gì?"
                ),
                chunks_by_id=by_id,
                relevant_chunk_ids=[
                    price["chunk_id"],
                    utility["chunk_id"],
                    plan_item["chunk_id"],
                ],
                relevant_chunk_groups=[
                    [price["chunk_id"]],
                    [utility["chunk_id"]],
                    [plan_item["chunk_id"]],
                ],
                target_type="multi_hop",
                difficulty="hard",
                required_metadata_fields=[
                    "document_type",
                    "year",
                    "project_name",
                    "project_status",
                ],
                must_include_terms=[price["price"], utility["amenity_term"], *plan_terms],
                evidence_fact_ids=[
                    price["fact_id"],
                    utility["fact_id"],
                    plan_item["fact_id"],
                ],
            )
        )

    project_names = list(dict.fromkeys(fact["project"] for fact in price_facts))
    for index in range(1, 31):
        year = 2022 if index <= 15 else 2027
        project = project_names[(index - 1) % len(project_names)]
        tests.append(
            _case(
                primary_slice="null_insufficient",
                index=index,
                query=(
                    f"Có market report giá nhà năm {year} trong bộ tài liệu "
                    f"cho dự án {project} không?"
                ),
                chunks_by_id=by_id,
                target_type="null",
                answerable=False,
                difficulty="hard",
                required_metadata_fields=["year", "document_type", "project_name"],
                metadata_conditions=[
                    _condition("year", year),
                    _condition("document_type", "market_report"),
                    _condition("project_name", project),
                ],
                evidence_fact_ids=[f"null_{year}_{index:02d}"],
            )
        )

    for index, table_case in enumerate(table_cases[:15], start=1):
        protected_id = table_case["relevant_chunk_ids"][0]
        protected_chunk = by_id[protected_id]
        pair_id = f"permission_pair_{index:02d}"
        shared_query = f"Phiên làm việc này cần tra cứu: {table_case['query']}"
        sensitive_terms = [table_case["expected"]["expected_cell_value"]]
        tests.append(
            _case(
                primary_slice="permission_sensitive",
                index=index,
                case_suffix="allow",
                query=shared_query,
                chunks_by_id=by_id,
                relevant_chunk_ids=[protected_id],
                target_type="permission_allowed",
                answerable=True,
                difficulty="hard",
                required_metadata_fields=[
                    "owner_id",
                    "notebook_id",
                    "document_ids",
                    "visibility",
                ],
                must_include_terms=list(table_case["expected"]["must_include_terms"]),
                query_context={
                    "owner_id": OWNER_ID,
                    "notebook_id": NOTEBOOK_ID,
                    "document_ids": [protected_chunk.document_id],
                    "groups": ["benchmark-reader"],
                },
                response_class="ANSWER_WITH_CITATIONS",
                permission_pair_id=pair_id,
                scenario_id=pair_id,
                evidence_fact_ids=list(table_case["evidence_fact_ids"]),
            )
        )
        tests.append(
            _case(
                primary_slice="permission_sensitive",
                index=index,
                case_suffix="deny",
                query=shared_query,
                chunks_by_id=by_id,
                protected_chunk_ids=[protected_id],
                target_type="permission_denied",
                answerable=False,
                difficulty="hard",
                required_metadata_fields=[
                    "owner_id",
                    "notebook_id",
                    "document_ids",
                    "visibility",
                ],
                must_include_terms=[],
                must_not_include_terms=sensitive_terms,
                must_not_cite_document_titles=[protected_chunk.document_title],
                query_context={
                    "owner_id": OWNER_ID,
                    "notebook_id": NOTEBOOK_ID,
                    "document_ids": [],
                    "groups": [],
                },
                response_class="ACCESS_DENIED_OR_INSUFFICIENT_AUTHORIZED_EVIDENCE",
                permission_pair_id=pair_id,
                scenario_id=pair_id,
                evidence_fact_ids=list(table_case["evidence_fact_ids"]),
            )
        )
    return tests


def _validate(tests: list[dict[str, Any]], corpus: list[EvalChunk]) -> list[dict[str, Any]]:
    if len(tests) != 300:
        raise ValueError(f"Expected exactly 300 queries, got {len(tests)}")
    primary_counts = Counter(test["primary_slice"] for test in tests)
    if set(primary_counts) != set(PRIMARY_SLICES):
        raise ValueError(f"Unexpected slices: {sorted(primary_counts)}")
    if any(primary_counts[name] != 30 for name in PRIMARY_SLICES):
        raise ValueError(f"Every primary slice must contain 30 cases: {primary_counts}")
    ids = [test["id"] for test in tests]
    if len(ids) != len(set(ids)):
        raise ValueError("Duplicate query IDs")
    by_query: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for test in tests:
        by_query[test["query"]].append(test)
    invalid_duplicates = []
    for query, duplicate_rows in by_query.items():
        if len(duplicate_rows) == 1:
            continue
        pair_ids = {row.get("permission_pair_id") for row in duplicate_rows}
        target_types = {row["target_type"] for row in duplicate_rows}
        contexts = {
            json.dumps(row["query_context"], ensure_ascii=False, sort_keys=True)
            for row in duplicate_rows
        }
        valid_pair = (
            len(duplicate_rows) == 2
            and len(pair_ids) == 1
            and None not in pair_ids
            and target_types == {"permission_allowed", "permission_denied"}
            and len(contexts) == 2
        )
        if not valid_pair:
            invalid_duplicates.append(query)
    if invalid_duplicates:
        raise ValueError(f"Invalid duplicate query text: {invalid_duplicates[:3]}")

    chunks_by_id = {chunk.id: chunk for chunk in corpus}
    audit: list[dict[str, Any]] = []
    for test in tests:
        referenced = {
            *test["relevant_chunk_ids"],
            *test["protected_chunk_ids"],
            *(chunk_id for group in test["relevant_chunk_groups"] for chunk_id in group),
        }
        missing_ids = sorted(referenced - set(chunks_by_id))
        if missing_ids:
            raise ValueError(f"{test['id']} references missing chunks: {missing_ids}")
        target_type = test["target_type"]
        if test["answerable"] and not test["relevant_chunk_ids"]:
            raise ValueError(f"{test['id']} is answerable but has no relevant chunks")
        if target_type == "multi_hop" and len(test["relevant_chunk_groups"]) < 2:
            raise ValueError(f"{test['id']} needs at least two evidence groups")
        if target_type == "null" and referenced:
            raise ValueError(f"{test['id']} null case must not reference evidence")
        if target_type == "permission_denied" and not test["protected_chunk_ids"]:
            raise ValueError(f"{test['id']} denied permission case has no protected evidence")
        if target_type == "permission_allowed" and not test["relevant_chunk_ids"]:
            raise ValueError(f"{test['id']} allowed permission case has no relevant evidence")
        if target_type == "null" and (
            test["retrieval_filters"].get("unsupported_field_policy") != "fail_closed"
        ):
            raise ValueError(f"{test['id']} null filter must be fail_closed")

        evidence_ids = (
            test["protected_chunk_ids"] if target_type == "permission_denied" else list(referenced)
        )
        evidence_text = "\n".join(chunks_by_id[chunk_id].text for chunk_id in evidence_ids)
        terms = test["expected"]["must_include_terms"]
        missing_terms = [term for term in terms if not _contains(evidence_text, term)]
        if missing_terms:
            raise ValueError(f"{test['id']} has terms absent from evidence: {missing_terms}")
        for term in terms:
            if re.search(r"\d$", term) and re.search(
                rf"{re.escape(term)}[.,]\d{{3}}", evidence_text
            ):
                raise ValueError(f"{test['id']} has a truncated numeric gold term: {term!r}")

        if target_type == "permission_denied":
            if terms:
                raise ValueError(f"{test['id']} denied case must not contain must_include_terms")
            sensitive_terms = test["expected"].get("must_not_include_terms", [])
            if not sensitive_terms:
                raise ValueError(f"{test['id']} denied case has no anti-leak terms")
            if any(not _contains(evidence_text, term) for term in sensitive_terms):
                raise ValueError(f"{test['id']} anti-leak terms are absent from protected evidence")

        if target_type == "multi_hop":
            required_titles = set(test["expected"].get("must_cite_document_titles", []))
            if required_titles != set(test["relevant_doc_titles"]):
                raise ValueError(f"{test['id']} does not require every relevant document citation")
            if len(test["relevant_doc_titles"]) > 1 and test["expected"]["document_title"]:
                raise ValueError(f"{test['id']} uses ambiguous singular document_title")

        if test["primary_slice"] == "table_structured":
            required_table_fields = {
                "table_id",
                "row_key",
                "column_key",
                "expected_cell_value",
                "header_path",
                "page",
                "source_block_ids",
                "bbox",
            }
            missing_table_fields = sorted(required_table_fields - set(test["expected"]))
            if missing_table_fields:
                raise ValueError(f"{test['id']} missing table fields: {missing_table_fields}")
            chunk = chunks_by_id[test["relevant_chunk_ids"][0]]
            if test["expected"]["table_id"] != chunk.table_identity:
                raise ValueError(f"{test['id']} table_id does not match parser provenance")
            if test["expected"]["source_block_ids"] != list(chunk.source_block_ids):
                raise ValueError(f"{test['id']} source_block_ids do not match parser provenance")
        audit.append(
            {
                "query_id": test["id"],
                "primary_slice": test["primary_slice"],
                "target_type": target_type,
                "status": "resolved" if not missing_ids else "unresolved",
                "relevant_chunk_count": len(test["relevant_chunk_ids"]),
                "evidence_group_count": len(test["relevant_chunk_groups"]),
                "protected_chunk_count": len(test["protected_chunk_ids"]),
                "must_term_count": len(terms),
                "must_term_coverage": 1.0
                if not terms
                else (len(terms) - len(missing_terms)) / len(terms),
                "relevant_documents": " | ".join(test["relevant_doc_titles"]),
                "relevant_chunk_ids": " | ".join(test["relevant_chunk_ids"]),
                "protected_chunk_ids": " | ".join(test["protected_chunk_ids"]),
                "query_style": test["query_style"],
                "filter_policy": test["retrieval_filters"]["unsupported_field_policy"],
            }
        )

    permission_pairs: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for test in tests:
        if test.get("permission_pair_id"):
            permission_pairs[str(test["permission_pair_id"])].append(test)
    if len(permission_pairs) != 15:
        raise ValueError(f"Expected 15 permission pairs, got {len(permission_pairs)}")
    for pair_id, pair in permission_pairs.items():
        if len(pair) != 2 or len({row["query"] for row in pair}) != 1:
            raise ValueError(f"{pair_id} is not a same-query allow/deny pair")

    content_rows = [test for test in tests if test["primary_slice"] == "content_only"]
    copied_gold_count = sum(
        any(_contains(test["query"], term) for term in test["expected"]["must_include_terms"])
        for test in content_rows
    )
    if copied_gold_count > 3:
        raise ValueError(f"Too many content-only queries copy gold terms: {copied_gold_count}")
    if set(Counter(test["query_style"] for test in tests)) != set(QUERY_STYLES):
        raise ValueError("Query style diversity is incomplete")
    return audit


def build_benchmark(
    source_dir: Path,
    *,
    approval: dict[str, Any] | None = None,
) -> tuple[
    list[dict[str, Any]],
    dict[str, Any],
    dict[str, Any],
    list[dict[str, Any]],
    list[EvalChunk],
]:
    missing = [
        filename for filename in QUERY_SOURCE_DOCUMENTS if not (source_dir / filename).exists()
    ]
    if missing:
        raise FileNotFoundError(f"Missing uploaded source document(s): {', '.join(missing)}")

    price_facts, summaries = _extract_price_facts(source_dir)
    utility_facts = _extract_utility_facts(source_dir)
    numeric_fact_audit = _numeric_fact_integrity_audit(utility_facts)
    numeric_fact_failures = [row for row in numeric_fact_audit if row["status"] != "match"]
    if numeric_fact_failures:
        raise ValueError(
            "Utility numeric anchors differ from source clauses: "
            f"{[row['fact_id'] for row in numeric_fact_failures]}"
        )
    statuses, milestones = _extract_plan_facts(source_dir)
    gold_config = _build_gold_config(price_facts, utility_facts, statuses, milestones)
    corpus, source_snapshots = build_corpus(
        source_dir=source_dir,
        gold_config=gold_config,
        chunk_size=600,
        chunk_overlap=80,
    )
    chunks = _attach_chunks(
        corpus,
        price_facts,
        summaries,
        utility_facts,
        statuses,
        milestones,
    )
    tests = _build_testset(
        chunks=chunks,
        price_facts=price_facts,
        summaries=summaries,
        utility_facts=utility_facts,
        statuses=statuses,
        milestones=milestones,
    )
    audit = _validate(tests, corpus)
    frozen = approval is not None
    if frozen:
        approved_on = str(approval.get("approved_on") or "")
        approval_basis = str(approval.get("approval_basis") or "")
        if not approved_on or not approval_basis:
            raise ValueError("Frozen approval requires approved_on and approval_basis")
        for test in tests:
            test["annotation_status"] = "approved_frozen_gold"
            test["human_review"] = {
                "status": "approved",
                "approved_on": approved_on,
                "reviewer_role": "benchmark_owner",
                "approval_basis": approval_basis,
            }
        gold_config["annotation_policy"] = (
            "Evidence is source-anchored, validated, and approved by the benchmark owner. "
            "Any content or source-hash change requires a new human review."
        )

    testset_sha256 = _jsonl_sha256(tests)
    query_source_hashes = {
        filename: _sha256(source_dir / filename) for filename in QUERY_SOURCE_DOCUMENTS
    }
    source_bundle_sha256 = _json_sha256(query_source_hashes)
    gold_metadata_sha256 = _json_sha256(gold_config)
    if approval and approval.get("testset_sha256") not in (None, "", testset_sha256):
        raise ValueError(
            "Frozen benchmark fingerprint changed. Do not reuse the prior approval; "
            "review the changed testset and create a new approval."
        )
    if approval and approval.get("source_bundle_sha256") not in (
        None,
        "",
        source_bundle_sha256,
    ):
        raise ValueError(
            "Frozen source-document fingerprint changed. Review the source changes before "
            "creating a new approval."
        )
    if approval and approval.get("gold_metadata_sha256") not in (
        None,
        "",
        gold_metadata_sha256,
    ):
        raise ValueError(
            "Frozen gold-metadata fingerprint changed. Review the metadata-rule changes "
            "before creating a new approval."
        )

    primary_counts = Counter(test["primary_slice"] for test in tests)
    all_slice_counts = Counter(
        slice_name for test in tests for slice_name in test["benchmark_slices"]
    )
    field_counts = Counter(field for test in tests for field in test["required_metadata_fields"])
    relevant_frequency = Counter(
        chunk_id for test in tests for chunk_id in test["relevant_chunk_ids"]
    )
    query_source_snapshots = [
        snapshot for snapshot in source_snapshots if snapshot["filename"] in QUERY_SOURCE_DOCUMENTS
    ]
    manifest = {
        "schema_version": "3.0",
        "name": "vinhomes_real_document_metadata_benchmark",
        "version": ("2026-08-04.real-v3-frozen" if frozen else "2026-08-04.real-v3-candidate"),
        "release_status": "frozen_gold" if frozen else "candidate",
        "benchmark_kind": "real_document_metadata_benchmark",
        "language": "vi",
        "query_count": len(tests),
        "retrieval_corpus_document_count": len(gold_config["documents"]),
        "query_source_document_count": len(QUERY_SOURCE_DOCUMENTS),
        "retrieval_corpus_chunk_count": len(corpus),
        "query_source_chunk_count": sum(
            snapshot["chunk_count"] for snapshot in query_source_snapshots
        ),
        "price_fact_count": len(price_facts),
        "utility_project_count": len(utility_facts),
        "project_status_count": len(statuses),
        "milestone_count": len(milestones),
        "numeric_fact_integrity": {
            "audited_fact_count": len(numeric_fact_audit),
            "thousands_separator_fact_count": sum(
                bool(row["contains_thousands_separator"]) for row in numeric_fact_audit
            ),
            "failed_count": len(numeric_fact_failures),
            "rows": numeric_fact_audit,
        },
        "primary_slice_counts": dict(sorted(primary_counts.items())),
        "all_slice_counts": dict(sorted(all_slice_counts.items())),
        "required_metadata_field_counts": dict(sorted(field_counts.items())),
        "split_counts": dict(sorted(Counter(test["split"] for test in tests).items())),
        "target_type_counts": dict(sorted(Counter(test["target_type"] for test in tests).items())),
        "query_style_counts": dict(sorted(Counter(test["query_style"] for test in tests).items())),
        "scenario_count": len({test["scenario_id"] for test in tests}),
        "evidence_fact_count": len(
            {fact_id for test in tests for fact_id in test["evidence_fact_ids"]}
        ),
        "ground_truth": {
            "status": "approved_frozen_gold"
            if frozen
            else "evidence_anchored_pending_human_review",
            "single": "exact deterministic chunk IDs",
            "multi_hop": "all relevant_chunk_groups are required",
            "null": "no matching chunk after supported metadata filters",
            "permission": (
                "15 same-query allow/deny pairs; only query_context access scope changes"
            ),
            "table_structured": "cell labels plus parser-native table provenance",
            "version_conflict": (
                "latest resolution, temporal qualification, and same-period scope differences"
            ),
            "unresolved_count": sum(row["status"] != "resolved" for row in audit),
        },
        "approval": (
            {
                **approval,
                "status": "approved_frozen_gold",
                "approved_case_count": len(tests),
                "testset_sha256": testset_sha256,
                "source_bundle_sha256": source_bundle_sha256,
                "gold_metadata_sha256": gold_metadata_sha256,
            }
            if approval
            else None
        ),
        "chunk_dependency": {
            "unique_relevant_chunk_count": len(relevant_frequency),
            "max_queries_per_relevant_chunk": max(relevant_frequency.values(), default=0),
            "note": (
                "Price tables are atomic under structure_recursive, so several row-level facts "
                "share one relevant chunk. Use clustered bootstrap and review this dependency."
            ),
        },
        "source_documents": source_snapshots,
        "query_source_sha256": query_source_hashes,
        "limitations": [
            (
                "Queries were human-approved, but future source, parser, chunking, or testset "
                "changes require a new review."
                if frozen
                else "Queries still require human relevance review before frozen-gold use."
            ),
            (
                "The table_structured slice covers DOCX tables; no image, chart, scan, "
                "or OCR source is present."
            ),
            (
                "Permission labels exercise the isolated harness ACL scope, not an "
                "external production ACL."
            ),
            (
                "Same-period scope-difference cases require an answer evaluator to "
                "judge qualification language."
            ),
            "DOCX pagination is unavailable, so ground truth uses chunk IDs and section headings.",
            "Hashing embeddings are only a free structural smoke test, not production quality.",
        ],
        "recommended_decision_metrics": {
            "answerable_recall_at_5_min": 0.85,
            "multi_hop_all_groups_at_10_min": 0.80,
            "null_rejection_at_10_min": 0.95,
            "permission_leak_at_10_max": 0.0,
            "permission_allowed_recall_at_10_min": 0.95,
            "sensitive_term_leak_at_10_max": 0.0,
            "table_structured_success_at_10_min": 0.90,
            "filter_preflight_pass_rate_min": 1.0,
        },
    }
    return tests, gold_config, manifest, audit, corpus


def _query_rows(tests: list[dict[str, Any]]) -> list[dict[str, Any]]:
    rows = []
    for test in tests:
        approved = test["annotation_status"] == "approved_frozen_gold"
        human_evidence_ids = [
            *test["relevant_chunk_ids"],
            *test["protected_chunk_ids"],
        ]
        review = test.get("human_review")
        review = review if isinstance(review, dict) else {}
        rows.append(
            {
                "query_id": test["id"],
                "split": test["split"],
                "scenario_id": test["scenario_id"],
                "query_style": test["query_style"],
                "permission_pair_id": test.get("permission_pair_id") or "",
                "primary_slice": test["primary_slice"],
                "benchmark_slices": " | ".join(test["benchmark_slices"]),
                "query": test["query"],
                "answerable": test["answerable"],
                "target_type": test["target_type"],
                "required_metadata_fields": " | ".join(test["required_metadata_fields"]),
                "relevant_documents": " | ".join(test["relevant_doc_titles"]),
                "relevant_chunk_ids": " | ".join(test["relevant_chunk_ids"]),
                "protected_chunk_ids": " | ".join(test["protected_chunk_ids"]),
                "must_include_terms": " | ".join(test["expected"]["must_include_terms"]),
                "must_not_include_terms": " | ".join(
                    test["expected"].get("must_not_include_terms", [])
                ),
                "must_cite_document_titles": " | ".join(
                    test["expected"].get("must_cite_document_titles", [])
                ),
                "expected_response_class": test["expected"].get("response_class", ""),
                "table_id": test["expected"].get("table_id", ""),
                "row_key": test["expected"].get("row_key", ""),
                "column_key": test["expected"].get("column_key", ""),
                "expected_cell_value": test["expected"].get("expected_cell_value", ""),
                "metadata_conditions": json.dumps(
                    test["retrieval_filters"]["metadata_conditions"], ensure_ascii=False
                ),
                "annotation_status": test["annotation_status"],
                "human_valid": "true" if approved else "",
                "human_relevant_chunk_ids": (
                    " | ".join(dict.fromkeys(human_evidence_ids)) if approved else ""
                ),
                "reviewer_notes": review.get("approval_basis", "") if approved else "",
            }
        )
    return rows


def write_benchmark(
    *,
    output_dir: Path,
    tests: list[dict[str, Any]],
    gold_config: dict[str, Any],
    manifest: dict[str, Any],
    audit: list[dict[str, Any]],
) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    _write_jsonl(output_dir / "testset.jsonl", tests)
    _write_json(output_dir / "gold_metadata.json", gold_config)
    _write_json(output_dir / "manifest.json", manifest)
    if manifest.get("approval"):
        _write_json(output_dir / "approval.json", manifest["approval"])
    _write_json(output_dir / "ablation_matrix.json", ABLATION_MATRIX)
    _write_csv(output_dir / "queries_for_review.csv", _query_rows(tests))
    _write_csv(output_dir / "ground_truth_audit.csv", audit)
    _write_csv(
        output_dir / "numeric_fact_integrity_audit.csv",
        manifest["numeric_fact_integrity"]["rows"],
    )
    primary = Counter(test["primary_slice"] for test in tests)
    all_slices = Counter(slice_name for test in tests for slice_name in test["benchmark_slices"])
    _write_csv(
        output_dir / "slice_distribution.csv",
        [
            {
                "slice": name,
                "primary_query_count": primary[name],
                "all_tagged_query_count": all_slices[name],
                "minimum_required": 30,
                "passes_minimum": primary[name] >= 30,
            }
            for name in PRIMARY_SLICES
        ],
    )


def _parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source-dir", type=Path, default=Path.home() / "Downloads")
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT_DIR)
    parser.add_argument(
        "--freeze",
        action="store_true",
        help="Record explicit human approval and freeze the testset fingerprint.",
    )
    return parser.parse_args()


def _existing_approval(output_dir: Path) -> dict[str, Any] | None:
    path = output_dir / "approval.json"
    if not path.exists():
        return None
    payload = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(payload, dict) or payload.get("status") != "approved_frozen_gold":
        raise ValueError(f"Invalid frozen approval file: {path}")
    return payload


def main() -> None:
    args = _parse_args()
    approval = _existing_approval(args.output_dir)
    if args.freeze and approval is None:
        approval = {
            "schema_version": "1.0",
            "status": "approved_frozen_gold",
            "approved_on": date.today().isoformat(),
            "reviewer_role": "benchmark_owner",
            "approval_basis": "Explicit benchmark-owner confirmation after human review.",
        }
    tests, gold_config, manifest, audit, _ = build_benchmark(
        args.source_dir,
        approval=approval,
    )
    write_benchmark(
        output_dir=args.output_dir,
        tests=tests,
        gold_config=gold_config,
        manifest=manifest,
        audit=audit,
    )
    print(f"Wrote {len(tests)} real-document queries to {args.output_dir / 'testset.jsonl'}")
    print(f"Wrote metadata gold config to {args.output_dir / 'gold_metadata.json'}")
    print(f"Wrote human review sheet to {args.output_dir / 'queries_for_review.csv'}")
    print(f"Ground-truth unresolved={manifest['ground_truth']['unresolved_count']}")
    print(f"Release status={manifest['release_status']}; testset_sha256={_jsonl_sha256(tests)}")


if __name__ == "__main__":
    main()
