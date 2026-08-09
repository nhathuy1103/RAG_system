"""Regression tests for authoritative identity across parser formats."""

from __future__ import annotations

import unicodedata
from hashlib import sha256
from io import BytesIO

from docx import Document as DocxDocument
from openpyxl import Workbook  # type: ignore[import-untyped]

from app.knowledge_quality.application.analysis import (
    is_auto_identity_eligible,
    strict_normalize_text,
)
from app.pipeline.documents.adapters.parsers import (
    CsvParser,
    DocxParser,
    HtmlParser,
    MarkdownParser,
    TxtParser,
    XlsxParser,
)
from app.pipeline.documents.application.content_identity import (
    build_parsed_document_fingerprint,
    build_template_structure_signature,
)
from app.pipeline.documents.domain.parsed import (
    ParsedDocument,
    ParsedImageMetadata,
    ParsedPage,
)


def _docx_prose() -> bytes:
    document = DocxDocument()
    document.add_heading("Quarterly Report", level=1)
    paragraph = document.add_paragraph()
    paragraph.add_run("Revenue").bold = True
    paragraph.add_run(" increased in 2026.")
    document.add_paragraph("Costs stayed controlled for all departments.")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _xlsx_table(rows: list[list[str]]) -> bytes:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Internal export sheet"
    for row in rows:
        worksheet.append(row)
    output = BytesIO()
    workbook.save(output)
    workbook.close()
    return output.getvalue()


def _docx_table(rows: list[list[str]]) -> bytes:
    document = DocxDocument()
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            table.cell(row_index, column_index).text = value
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_prose_identity_matches_across_txt_docx_markdown_and_html() -> None:
    sources = {
        "txt": (
            b"Quarterly Report\nRevenue increased in 2026.\n"
            b"Costs stayed controlled for all departments."
        ),
        "docx": _docx_prose(),
        "markdown": (
            b"# Quarterly Report\n\n**Revenue** increased in 2026.\n\n"
            b"Costs stayed controlled for all departments."
        ),
        "html": (
            b"<h1>Quarterly Report</h1><p><strong>Revenue</strong> increased in 2026.</p>"
            b"<p>Costs stayed controlled for all departments.</p>"
        ),
    }
    parsed = {
        "txt": TxtParser().parse(sources["txt"]),
        "docx": DocxParser().parse(sources["docx"]),
        "markdown": MarkdownParser().parse(sources["markdown"]),
        "html": HtmlParser().parse(sources["html"]),
    }

    fingerprints = {
        name: build_parsed_document_fingerprint(document) for name, document in parsed.items()
    }

    assert len({sha256(content).hexdigest() for content in sources.values()}) == 4
    assert len({fingerprint.strict_hash for fingerprint in fingerprints.values()}) == 1
    assert all(is_auto_identity_eligible(fingerprint) for fingerprint in fingerprints.values())


def test_structured_table_identity_matches_across_csv_xlsx_docx_and_html() -> None:
    rows = [
        ["Metric", "Annual value"],
        ["Revenue", "120 million USD"],
        ["Operating cost", "80 million USD"],
    ]
    csv_content = b"Metric,Annual value\nRevenue,120 million USD\nOperating cost,80 million USD\n"
    html_content = (
        b"<table><tr><th>Metric</th><th>Annual value</th></tr>"
        b"<tr><td>Revenue</td><td>120 million USD</td></tr>"
        b"<tr><td>Operating cost</td><td>80 million USD</td></tr></table>"
    )
    parsed = (
        CsvParser().parse(csv_content),
        XlsxParser().parse(_xlsx_table(rows)),
        DocxParser().parse(_docx_table(rows)),
        HtmlParser().parse(html_content),
    )

    fingerprints = tuple(build_parsed_document_fingerprint(item) for item in parsed)

    assert len({fingerprint.strict_hash for fingerprint in fingerprints}) == 1
    assert {fingerprint.table_count for fingerprint in fingerprints} == {1}
    assert all(is_auto_identity_eligible(fingerprint) for fingerprint in fingerprints)


def test_table_value_or_order_change_does_not_match() -> None:
    original = CsvParser().parse(
        b"Metric,Annual value\nRevenue,120 million USD\nOperating cost,80 million USD\n"
    )
    changed = CsvParser().parse(
        b"Metric,Annual value\nRevenue,121 million USD\nOperating cost,80 million USD\n"
    )
    reordered = CsvParser().parse(
        b"Metric,Annual value\nOperating cost,80 million USD\nRevenue,120 million USD\n"
    )

    hashes = {
        build_parsed_document_fingerprint(item).strict_hash
        for item in (original, changed, reordered)
    }

    assert len(hashes) == 3
    assert build_template_structure_signature(original) == build_template_structure_signature(
        changed
    )


def test_template_structure_signature_changes_with_business_columns() -> None:
    original = CsvParser().parse(b"Unit,List price\nA101,4500000000\n")
    changed_schema = CsvParser().parse(b"Unit,Discounted price\nA101,4200000000\n")

    assert build_template_structure_signature(original) != build_template_structure_signature(
        changed_schema
    )


def test_flat_table_text_is_not_auto_equated_with_proven_structured_cells() -> None:
    flat = TxtParser().parse(
        b"Metric Annual value Revenue 120 million USD Operating cost 80 million USD"
    )
    structured = CsvParser().parse(
        b"Metric,Annual value\nRevenue,120 million USD\nOperating cost,80 million USD\n"
    )

    flat_fingerprint = build_parsed_document_fingerprint(flat)
    structured_fingerprint = build_parsed_document_fingerprint(structured)

    assert flat_fingerprint.strict_hash != structured_fingerprint.strict_hash
    assert flat_fingerprint.table_count == 0
    assert structured_fingerprint.table_count == 1


def test_authoritative_normalization_preserves_semantic_unicode() -> None:
    decomposed = unicodedata.normalize("NFD", "Báo cáo quý")

    assert strict_normalize_text(decomposed) == "Báo cáo quý"
    assert strict_normalize_text("x²") != strict_normalize_text("x2")
    assert strict_normalize_text("👩\u200d💻") != strict_normalize_text("👩💻")


def test_unrepresented_visual_content_blocks_automatic_identity() -> None:
    parsed = ParsedDocument(
        text="Quarterly report with a chart and enough explanatory text for identity.",
        pages=[
            ParsedPage(
                page_number=1,
                text="Quarterly report with a chart and enough explanatory text for identity.",
            )
        ],
        images_metadata=[ParsedImageMetadata(location="page:1:image:1")],
        parser_name="pdf",
    )

    fingerprint = build_parsed_document_fingerprint(parsed)

    assert fingerprint.unrepresented_visual_count == 1
    assert fingerprint.identity_trusted is False
    assert is_auto_identity_eligible(fingerprint) is False
