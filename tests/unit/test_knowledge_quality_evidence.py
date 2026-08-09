"""Tests for reviewer-facing relation evidence assembly."""

from datetime import UTC, datetime
from io import BytesIO
from uuid import UUID

from docx import Document as DocxDocument

from app.knowledge_quality.application.evidence import build_relation_chunk_pairs
from app.knowledge_quality.application.original_preview import build_original_review_blocks
from app.knowledge_quality.domain.models import (
    DocumentRelation,
    RelationEvidenceChunk,
    RelationEvidenceChunkPair,
    RelationEvidenceDocument,
    RelationStatus,
    RelationType,
)

OWNER_ID = UUID("20000000-0000-0000-0000-000000000002")
NOTEBOOK_ID = UUID("10000000-0000-0000-0000-000000000001")
SOURCE_ID = UUID("30000000-0000-0000-0000-000000000003")
TARGET_ID = UUID("40000000-0000-0000-0000-000000000004")
NOW = datetime(2026, 7, 30, tzinfo=UTC)


def _relation() -> DocumentRelation:
    return DocumentRelation(
        id=UUID("50000000-0000-0000-0000-000000000005"),
        owner_id=OWNER_ID,
        notebook_id=NOTEBOOK_ID,
        source_document_id=SOURCE_ID,
        target_document_id=TARGET_ID,
        relation_type=RelationType.CONFLICT_CANDIDATE,
        status=RelationStatus.PENDING,
        confidence=0.88,
        signals={
            "selected_chunk_pair": {
                "source_chunk_index": 1,
                "target_chunk_index": 1,
            }
        },
        reason="date_mismatch",
        detector_version="knowledge-quality-v2",
        preferred_document_id=None,
        resolved_by=None,
        resolved_at=None,
        created_at=NOW,
        updated_at=NOW,
    )


def _chunk(
    document_id: UUID,
    index: int,
    content: str,
    *,
    group_id: str | None = None,
) -> RelationEvidenceChunk:
    return RelationEvidenceChunk(
        id=f"{document_id}:{index}",
        document_id=document_id,
        chunk_index=index,
        content=content,
        page_number=1,
        section_title=None,
        normalized_content_hash=None,
        exact_duplicate_group_id=group_id,
    )


def test_relation_evidence_groups_exact_conflict_and_added_chunks() -> None:
    source_chunks = (
        _chunk(SOURCE_ID, 0, "Shared introduction", group_id="intro"),
        _chunk(SOURCE_ID, 1, "Policy effective date is 15/03/2027."),
        _chunk(SOURCE_ID, 2, "New approval step must be reviewed by CSKH."),
    )
    target_chunks = (
        _chunk(TARGET_ID, 0, "Shared introduction", group_id="intro"),
        _chunk(TARGET_ID, 1, "Policy effective date is 15/03/2026."),
    )

    pairs = build_relation_chunk_pairs(_relation(), source_chunks, target_chunks)

    assert [pair.evidence_type for pair in pairs] == [
        "exact_content",
        "conflict_candidate",
        "source_only",
    ]
    assert pairs[1].signals["date_agreement"] is False
    assert pairs[2].source_chunk is source_chunks[2]
    assert pairs[2].signals == {"added_content": True}


def test_original_docx_preview_highlights_changed_table_row() -> None:
    source_chunk = _chunk(
        SOURCE_ID,
        2,
        "Ma tai lieu | CSKH-RET-01\nNgay ban hanh | 15/03/2027\nPham vi | Online",
    )
    target_chunk = _chunk(
        TARGET_ID,
        2,
        "Ma tai lieu | CSKH-RET-01\nNgay ban hanh | 15/03/2026\nPham vi | Online",
    )
    pairs = (
        RelationEvidenceChunkPair(
            source_chunk=source_chunk,
            target_chunk=target_chunk,
            evidence_type="conflict_candidate",
            confidence=0.93,
            signals={"reason_codes": ["date_mismatch"]},
            reason="date_mismatch",
        ),
    )

    blocks = build_original_review_blocks(
        RelationEvidenceDocument(
            id=SOURCE_ID,
            original_filename="policy.docx",
            quality_status="review_required",
            version_number=1,
            is_current=True,
            canonical_document_id=None,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
        _docx_with_policy_table("15/03/2027"),
        pairs,
        side="source",
    )

    date_block = next(block for block in blocks if "15/03/2027" in block.text)
    unchanged_block = next(block for block in blocks if "CSKH-RET-01" in block.text)
    assert date_block.block_type == "table_row"
    assert date_block.cells == ("Ngay ban hanh", "15/03/2027")
    assert date_block.highlight_type == "conflict_candidate"
    assert date_block.matched_pair_index == 0
    assert "|" not in date_block.text
    assert unchanged_block.highlight_type is None


def _docx_with_policy_table(effective_date: str) -> bytes:
    document = DocxDocument()
    document.add_paragraph("Tai lieu demo")
    table = document.add_table(rows=3, cols=2)
    rows = [
        ("Ma tai lieu", "CSKH-RET-01"),
        ("Ngay ban hanh", effective_date),
        ("Pham vi", "Online"),
    ]
    for row, values in zip(table.rows, rows, strict=True):
        for cell, value in zip(row.cells, values, strict=True):
            cell.text = value
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
