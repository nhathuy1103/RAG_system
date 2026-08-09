from __future__ import annotations

from io import BytesIO

from docx import Document as DocxDocument

from app.pipeline.documents.adapters.parsers import CsvParser, DocxParser
from app.pipeline.documents.application.extraction_pipeline import sanitize_parsed_document
from app.pipeline.documents.domain.parsed import ParsedDocument, ParsedPage, ParsedSection
from app.pipeline.indexing.application.chunker import Chunker
from app.pipeline.shared.markdown import render_parsed_document_markdown


def _docx_with_heading_and_table() -> bytes:
    document = DocxDocument()
    document.add_heading("Travel policy", level=1)
    document.add_paragraph("Allowance applies to approved trips.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Role"
    table.cell(0, 1).text = "Limit"
    table.cell(1, 0).text = "Employee"
    table.cell(1, 1).text = "120 USD"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_docx_is_normalized_to_structured_markdown_before_chunking() -> None:
    parsed = DocxParser().parse(_docx_with_heading_and_table())

    assert parsed.document_metadata["content_format"] == "markdown"
    assert parsed.document_metadata["content_representation_version"] == "markdown-v1"
    assert parsed.content_markdown == (
        "# Travel policy\n\n"
        "Allowance applies to approved trips.\n\n"
        "| Role | Limit |\n"
        "| --- | --- |\n"
        "| Employee | 120 USD |"
    )

    chunks = Chunker.structure_recursive(chunk_size=64, overlap=8).chunk("doc-1", 1, parsed)

    assert chunks[0].text.startswith("# Travel policy")
    assert chunks[1].text.startswith("| Role | Limit |\n| --- | --- |")


def test_csv_table_is_normalized_to_markdown_table() -> None:
    parsed = CsvParser().parse(b"Role,Limit\nEmployee,120 USD\n")

    assert parsed.content_markdown == ("| Role | Limit |\n| --- | --- |\n| Employee | 120 USD |")


def test_pdf_pages_receive_markdown_page_sections() -> None:
    parsed = ParsedDocument(
        text="First page\n\nSecond page",
        pages=[
            ParsedPage(page_number=1, text="First page"),
            ParsedPage(page_number=2, text="Second page"),
        ],
        sections=[
            ParsedSection(text="First page", page_number=1, title="Page 1"),
            ParsedSection(text="Second page", page_number=2, title="Page 2"),
        ],
        parser_name="pdf",
    )

    assert render_parsed_document_markdown(parsed) == (
        "## Page 1\n\nFirst page\n\n## Page 2\n\nSecond page"
    )


def test_sanitization_rebuilds_stale_markdown_representation() -> None:
    parsed = ParsedDocument(
        text="Current content",
        pages=[ParsedPage(page_number=1, text="Current content")],
        sections=[ParsedSection(text="Current content", page_number=1, title="Text")],
        parser_name="txt",
        content_markdown="stale content",
    )

    sanitized = sanitize_parsed_document(parsed)

    assert sanitized.content_markdown == "Current content"
    assert sanitized.document_metadata["content_format"] == "markdown"
    assert sanitized.document_metadata["content_representation_version"] == "markdown-v1"
